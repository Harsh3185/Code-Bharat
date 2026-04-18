from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
BACKEND = ROOT / "Backend"
COMPILER = ROOT / "Compiler"
OUT = ROOT / "output" / "doc" / "Code-Bharat-Backend-Interview-Guide.docx"


def read(path):
    return path.read_text(encoding="utf-8")


def style(doc):
    s = doc.sections[0]
    s.top_margin = Inches(0.7)
    s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8)
    s.right_margin = Inches(0.8)
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)


def title(doc, text, size=22, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = not italic
    r.italic = italic
    r.font.size = Pt(size)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def nums(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8)


def numbered(text):
    lines = text.splitlines()
    width = len(str(max(len(lines), 1)))
    return "\n".join(f"{i:>{width}}: {line}" for i, line in enumerate(lines or [""], 1))


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            c[i].text = str(v)


def ranges(text, anchors):
    lines = text.splitlines()
    found = []
    for anchor, note in anchors:
        pos = None
        for i, line in enumerate(lines, 1):
            if anchor in line:
                pos = i
                break
        if pos is None:
            raise ValueError(f"Missing anchor: {anchor}")
        found.append((pos, note))
    out = []
    for i, (start, note) in enumerate(found):
        end = found[i + 1][0] - 1 if i + 1 < len(found) else len(lines)
        out.append((start, end, note))
    return out


def walkthrough(doc, text, anchors):
    for start, end, note in ranges(text, anchors):
        p = doc.add_paragraph()
        label = f"Lines {start}-{end}" if start != end else f"Line {start}"
        p.add_run(f"{label}: ").bold = True
        p.add_run(note)


def tree():
    files = sorted(
        p.relative_to(BACKEND).as_posix()
        for p in BACKEND.rglob("*")
        if p.is_file() and "node_modules" not in p.parts
    )
    out = ["Backend/"]
    for f in files:
        out.append(f"{'  ' * f.count('/')} - {f.split('/')[-1]}")
    return "\n".join(out)


ROUTES = [
    ("POST", "/api/auth/register", "Public", "checkSchema(registerValidatingSchema) -> register"),
    ("POST", "/api/auth/login", "Public", "checkSchema(loginValidatingSchema) -> login"),
    ("GET", "/api/auth/me", "Authenticated", "requireAuth -> me"),
    ("POST", "/api/auth/logout", "Cookie client", "logout"),
    ("GET", "/api/problems", "Public", "getProblemSet"),
    ("GET", "/api/problem/:Id", "Authenticated", "requireAuth -> getProblem"),
    ("POST", "/api/problem", "Admin", "requireAuth -> isAdmin -> checkSchema(problemValidatingSchema) -> addProblem"),
    ("DELETE", "/api/problem/:Id", "Admin", "requireAuth -> isAdmin -> deleteProblem"),
    ("POST", "/api/testcases/:id", "Admin", "requireAuth -> isAdmin -> checkSchema(testCaseValidatingSchema) -> addTestCase"),
    ("GET", "/api/testcases/:id", "Admin", "requireAuth -> isAdmin -> getTestCasesByProblem"),
    ("GET", "/api/leaderboard", "Public", "getLeaderboard"),
    ("GET", "/api/user/:id", "Public", "userProfile"),
    ("GET", "/api/profile", "Authenticated", "requireAuth -> getMyProfile"),
    ("PUT", "/api/profile", "Authenticated", "requireAuth -> checkSchema(profileValidatingSchema) -> editProfile"),
    ("GET", "/api/submissions", "Authenticated", "requireAuth -> showAllSubmissions"),
    ("GET", "/api/problem/:id/submissions", "Authenticated", "requireAuth -> showProblemSubmissions"),
    ("POST", "/api/problem/:id/submit", "Authenticated", "requireAuth -> checkSchema(submissionValidatingSchema) -> submitSolution"),
]


ENV_ROWS = [
    ("PORT", "Listen port"),
    ("MONGODB_URL", "MongoDB connection string"),
    ("COOKIE_SECRET", "Secret passed to cookie-parser"),
    ("JWT_SECRET", "JWT signing and verification secret"),
    ("ADMIN_SECRET", "Secret needed for admin registration"),
    ("COMPILER_SERVICE_URL", "Submission judge service URL"),
    ("NODE_ENV", "Controls secure cookie options"),
    ("CORS_ALLOWED_ORIGINS", "Optional comma-separated CORS allow-list override"),
    ("SESSION_SECRET", "Defined in .env but unused in source"),
]


GUIDES = []
COMPILER_GUIDES = []

GUIDES.extend([
    {
        "path": "package.json",
        "title": "Backend Package Manifest",
        "purpose": "Defines Node version, runtime scripts, package mode, and core dependencies.",
        "flow": "This file shapes how the backend boots and what libraries are available across the app.",
        "points": [
            "Real startup uses `node Index.mjs` from the scripts section.",
            "`.mjs` files still run as ES modules even though package type is `commonjs`.",
            "Pinned `node: 20.x` helps deployment consistency."
        ],
        "notes": [
            "The `main` field is stale because the app really starts from `Index.mjs`.",
            "No lint or test scripts are defined yet."
        ],
        "anchors": [
            ('"name": "backend"', "Package metadata starts here."),
            ('"engines": {', "Pins the Node.js runtime version."),
            ('"scripts": {', "Defines start and dev commands."),
            ('"type": "commonjs"', "Sets default package mode, although `.mjs` files override it."),
            ('"dependencies": {', "Declares the backend stack: Express, Mongoose, bcrypt, JWT, cookies, validation, axios, CORS, dotenv."),
        ],
    },
    {
        "path": ".gitignore",
        "title": "Git Ignore Rules",
        "purpose": "Prevents secrets, dependencies, logs, and local editor files from being committed.",
        "flow": "Not part of request execution, but critical for secure repository hygiene.",
        "points": [
            "Ignoring `.env` prevents secret leakage.",
            "Ignoring `node_modules` and local editor folders keeps the repo clean."
        ],
        "notes": ["`package-lock.json` is ignored here even though a lockfile exists in the repo."],
        "anchors": [
            ('node_modules', "Skips installed dependencies."),
            ('.env', "Skips secret environment files."),
            ('logs/', "Skips runtime log files."),
            ('coverage/', "Skips generated coverage and build output."),
            ('.vscode/', "Skips editor-specific files."),
        ],
    },
    {
        "path": ".dockerignore",
        "title": "Docker Ignore Rules",
        "purpose": "Keeps the Docker build context small and secret-safe.",
        "flow": "Affects container build behavior rather than live request handling.",
        "points": [
            "Skipping `node_modules` avoids copying host-specific packages into the image.",
            "Skipping `.env` keeps secrets out of Docker layers."
        ],
        "notes": ["Very small but high-value Docker ignore file."],
        "anchors": [
            ('node_modules', "Prevents local modules from entering the image context."),
            ('.env', "Prevents environment secrets from entering the image context."),
        ],
    },
    {
        "path": "Dockerfile",
        "title": "Backend Dockerfile",
        "purpose": "Builds a runnable container image for the backend.",
        "flow": "Defines the environment in which the Express server will run in containerized deployments.",
        "points": [
            "Uses Node 20, matching `package.json` engines.",
            "Copies package files before source to improve Docker layer caching.",
            "Starts the app with `node Index.mjs`."
        ],
        "notes": ["Could be improved with `npm ci`, a non-root user, and a slimmer production image."],
        "anchors": [
            ('FROM node:20', "Selects the Node 20 base image."),
            ('WORKDIR /app', "Sets the working directory inside the container."),
            ('COPY package*.json ./', "Copies package metadata first for cached installs."),
            ('RUN npm install', "Installs backend dependencies inside the image."),
            ('COPY . .', "Copies application source code."),
            ('EXPOSE 8000', "Documents the app port."),
            ('ENV NODE_ENV=production', "Marks the runtime as production."),
            ('CMD ["node", "Index.mjs"]', "Starts the backend server."),
        ],
    },
    {
        "path": "Index.mjs",
        "title": "Application Entrypoint",
        "purpose": "Bootstraps the server, database connection, CORS, parsers, routers, health route, 404 fallback, and listener.",
        "flow": "Every request first enters here, passes global middleware, then reaches a router/controller or the 404 handler.",
        "points": [
            "Global middleware order is important: CORS, body parsing, cookie parsing, then routers.",
            "The app uses cookie-based auth, so CORS is configured with credentials.",
            "Routers include full API paths and are mounted without a shared prefix."
        ],
        "notes": [
            "There is no central error middleware.",
            "Cookie parser receives a secret, but the auth cookie is not stored as a signed cookie."
        ],
        "anchors": [
            ("import express from 'express';", "Imports the server framework and all route/config modules needed to boot the app."),
            ('dotenv.config();', "Loads environment variables from `.env`."),
            ('const app = express();', "Creates the Express application instance."),
            ('DBConnection();', "Starts MongoDB connection at boot time."),
            ('const defaultAllowedOrigins = [', "Defines the hardcoded frontend allow list."),
            ('const allowedOrigins = process.env.CORS_ALLOWED_ORIGINS', "Allows deployments to override the CORS list with an env var."),
            ('const isAllowedOrigin = (origin) => {', "Adds custom logic for allowed origins, including Render preview domains."),
            ('app.use(cors({', "Registers credential-aware CORS before any route logic."),
            ('app.use(express.json());', "Registers JSON parser, URL-encoded parser, and cookie parser."),
            ('app.use(authRouter);', "Mounts each feature router."),
            ('app.get("/", (req, res) => {', "Defines a basic health-check endpoint."),
            ('app.use((req, res) => {', "Defines a catch-all 404 handler."),
            ('const PORT = process.env.PORT || 8000;', "Chooses the port and starts the listener."),
        ],
    },
    {
        "path": "src/Database/Database.mjs",
        "title": "Database Connection Module",
        "purpose": "Connects Mongoose to MongoDB.",
        "flow": "Runs during startup before routes begin handling database-backed requests.",
        "points": [
            "Database connection is wrapped in a dedicated function for cleaner bootstrapping.",
            "Mongoose is the abstraction layer over MongoDB."
        ],
        "notes": ["Connection errors are logged but not re-thrown, so the server can continue running even if MongoDB is unavailable."],
        "anchors": [
            ("import mongoose from 'mongoose';", "Imports Mongoose and dotenv."),
            ('dotenv.config();', "Loads environment variables for the database connection."),
            ('const DBConnection = async () => {', "Defines the async connection function."),
            ('const MONGODB_URL = process.env.MONGODB_URL;', "Reads the MongoDB connection string."),
            ('await mongoose.connect(MONGODB_URL);', "Opens the database connection."),
            ('console.log("Connected to the Database Successfully,")', "Logs successful connection."),
            ('console.log("Error connecting to database "', "Logs connection failure."),
            ('export default DBConnection;', "Exports the function for `Index.mjs`."),
        ],
    },
    {
        "path": "src/Utils/jwt.mjs",
        "title": "JWT Utility",
        "purpose": "Centralizes token creation and verification for cookie-based authentication.",
        "flow": "Used by login to issue tokens and by auth middleware to verify them on protected requests.",
        "points": [
            "Token payload contains identity and role data.",
            "Token expiry is aligned with cookie lifetime at two hours."
        ],
        "notes": [
            "Bug: `dotenv.config;` is missing parentheses.",
            "Bug: payload uses `user.name`, but the schema field is `userName`."
        ],
        "anchors": [
            ('import jwt from "jsonwebtoken"', "Imports JWT signing support and dotenv."),
            ('dotenv.config;', "Attempts to load env vars, but this line does not actually call the function."),
            ('const JWT_SECRET = process.env.JWT_SECRET;', "Reads the JWT signing secret."),
            ('export const createToken = (user) => {', "Builds a signed JWT from user details."),
            ("{ expiresIn: '2h' }", "Sets token expiration."),
            ('export const verifyToken = (token) => {', "Verifies a token and returns decoded payload."),
        ],
    },
    {
        "path": "src/Middleware/requireAuth.mjs",
        "title": "Authentication Middleware",
        "purpose": "Protects routes by reading the `token` cookie, verifying it, loading the user, and attaching `req.user`.",
        "flow": "Runs before protected controllers such as `me`, `getProblem`, `getMyProfile`, and submission endpoints.",
        "points": [
            "Middleware validates the cookie token and also re-reads the user from MongoDB.",
            "Controllers reuse `req.user` instead of re-querying identity again."
        ],
        "notes": ["This depends on frontend requests being sent with cookies and `credentials: true`."],
        "anchors": [
            ('import { verifyToken } from "../Utils/jwt.mjs";', "Imports token verification and the User model."),
            ('const requireAuth = async (req, res, next) => {', "Defines the route-protection middleware."),
            ('const token = req.cookies.token;', "Reads the auth token from cookies."),
            ('if (!token) {', "Rejects requests when no token is present."),
            ('const decoded = verifyToken(token);', "Decodes and verifies the JWT."),
            ("const user = await User.findById(decoded.id).select('-password');", "Loads the user and hides the password field."),
            ('req.user = user;', "Attaches the authenticated user to the request."),
            ('return res.status(401).json({ message: "Invalid or expired token" });', "Handles invalid or expired tokens."),
        ],
    },
    {
        "path": "src/Middleware/isAdmin.mjs",
        "title": "Admin Authorization Middleware",
        "purpose": "Allows only admin users to proceed to sensitive routes.",
        "flow": "Runs after `requireAuth` on admin-only routes like problem creation and test-case management.",
        "points": [
            "Authentication proves identity; authorization checks permission.",
            "Role comparison is lowercase, matching the model's pre-save normalization."
        ],
        "notes": ["If middleware order changes and `req.user` is missing, this denies access."],
        "anchors": [
            ('const isAdmin = (req, res, next) => {', "Defines the admin guard."),
            ("if (req.user && req.user.role.toLowerCase() === 'admin') {", "Allows the request only for authenticated admins."),
            ('return res.status(403).json({ message: "Admin access only" });', "Rejects non-admin users."),
            ('export default isAdmin;', "Exports the middleware."),
        ],
    },
])

GUIDES.extend([
    {
        "path": "src/Controllers/authControllers.mjs",
        "title": "Auth Controllers",
        "purpose": "Implements register, login, logout, and current-user responses.",
        "flow": "These functions run after route validation and, for `me`, after auth middleware has already attached `req.user`.",
        "points": [
            "Register hashes passwords before storage.",
            "Login reuses `req._foundUser` from validator to avoid a duplicate lookup.",
            "Auth state is stored in an HTTP-only cookie."
        ],
        "notes": ["Login returns 201 instead of 200.", "Cookie settings change by `NODE_ENV` to support cross-site production cookies."],
        "anchors": [
            ('const getCookieOptions = () => {', "Builds cookie settings for login."),
            ('const getClearCookieOptions = () => {', "Builds matching settings for logout cookie clearing."),
            ('export const register = async (req, res) => {', "Starts registration flow."),
            ('const errors = validationResult(req);', "Stops if route validation failed."),
            ('const data = matchedData(req, { includeOptionals: true });', "Uses only validated request fields."),
            ('const hashedPassword = await bcrypt.hash(password, 10);', "Hashes the password before storing it."),
            ('const user = await User.create({', "Creates the new user document."),
            ('const { password: _password, ...publicUser } = user.toObject();', "Removes password hash from API response."),
            ('export const login = async (req, res) => {', "Starts login flow."),
            ('const user = req._foundUser;', "Reuses user loaded by login validation."),
            ("res.cookie('token', token, getCookieOptions());", "Writes JWT into an HTTP-only cookie."),
            ('export const me = async (req, res) => {', "Returns authenticated user profile data."),
            ('export const logout = async (req, res) => {', "Clears the auth cookie."),
        ],
    },
    {
        "path": "src/Controllers/problemControllers.mjs",
        "title": "Problem Controllers",
        "purpose": "Handles reading, creating, listing, and deleting problem documents.",
        "flow": "Controllers receive validated/authenticated requests from problem routes and then perform Mongoose operations.",
        "points": [
            "ObjectId validation happens before database access.",
            "Create route records `created_by` from the authenticated admin.",
            "Delete route manually cleans related test cases and submissions."
        ],
        "notes": ["Problem list uses field projection for lighter responses."],
        "anchors": [
            ('export const getProblem = async (req, res) => {', "Starts single-problem read logic."),
            ('if (!mongoose.Types.ObjectId.isValid(Id)) {', "Rejects malformed ids early."),
            ('const problem = await Problem.findById(Id);', "Loads one problem document."),
            ('export const addProblem = async (req, res) => {', "Starts create-problem flow."),
            ('const data = matchedData(req, { includeOptionals: true });', "Extracts validated body fields."),
            ('const created_by = req.user.id;', "Uses authenticated admin id as creator."),
            ('const problem = await Problem.create({', "Creates the problem."),
            ('export const getProblemSet = async (req, res) => {', "Starts lightweight list flow."),
            ("const problems = await Problem.find().select('problemNumber title acceptanceRate');", "Fetches only list-page fields."),
            ('export const deleteProblem = async (req, res) => {', "Starts delete flow."),
            ('const deletedProblem = await Problem.findByIdAndDelete(Id);', "Deletes main problem document."),
            ('await TestCase.deleteMany({ problemId: Id });', "Deletes related hidden test cases."),
            ('await Submission.deleteMany({ problemId: Id });', "Deletes related submission history."),
        ],
    },
    {
        "path": "src/Controllers/profileControllers.mjs",
        "title": "Profile Controllers",
        "purpose": "Handles public profile reads, self-profile reads, leaderboard generation, and profile editing.",
        "flow": "Reads user documents directly or, for self-profile routes, uses authenticated identity from middleware.",
        "points": [
            "Leaderboard is computed in application code after fetching user data with `.lean()`.",
            "Profile edit updates only selected fields."
        ],
        "notes": ["Validator includes `firstName` and `lastName`, but controller/model do not use them."],
        "anchors": [
            ('export const userProfile = async (req, res) => {', "Starts public profile lookup."),
            ('const user = await User.findById(id).select(', "Fetches selected public profile fields."),
            ('export const getMyProfile = async (req, res) => {', "Starts authenticated self-profile lookup."),
            ('const user = await User.findById(req.user.id).select(', "Fetches selected current-user fields."),
            ('export const getLeaderboard = async (_req, res) => {', "Starts leaderboard generation."),
            ('const users = await User.find()', "Loads users needed for ranking."),
            ('const leaderboard = users', "Builds, sorts, and ranks the leaderboard."),
            ('export const editProfile = async (req, res) => {', "Starts profile update flow."),
            ('const { profilePicture, bio, location, institution } = req.body;', "Reads editable fields from request body."),
            ('await user.save();', "Persists profile changes."),
        ],
    },
    {
        "path": "src/Controllers/submissionControllers.mjs",
        "title": "Submission Controllers",
        "purpose": "Implements submission history and the full online-judge flow.",
        "flow": "This is the deepest request path: authenticated request -> problem lookup -> test-case lookup -> compiler calls -> comparison -> stats updates -> submission write -> response.",
        "points": [
            "Good interview example of orchestration across middleware, models, and an external service.",
            "Whitespace normalization reduces false-negative comparisons.",
            "Writes problem stats, user stats, and submission history in one request."
        ],
        "notes": [
            "History endpoints do not use `try/catch`.",
            "Stats updates are not transactional.",
            "Storing all judge inputs/outputs in submissions helps debugging but increases storage."
        ],
        "anchors": [
            ('const normalize = (s) => s.trim().split("\\n").map((l) => l.trim()).join("\\n");', "Defines output normalization for comparison."),
            ('export const showAllSubmissions = async (req, res) => {', "Starts all-submissions history route."),
            ('const submissions = await Submission.find({ userId: req.user._id })', "Loads current user's submission history."),
            ('export const showProblemSubmissions = async (req, res) => {', "Starts per-problem history route."),
            ('const submissions = await Submission.find({', "Filters history by user and problem."),
            ('export const submitSolution = async (req, res) => {', "Starts the main judge workflow."),
            ('const errors = validationResult(req);', "Rejects invalid input before expensive work."),
            ('const problemId = req.params.id;', "Reads the target problem id."),
            ('const { code, language } = matchedData(req);', "Extracts validated code and language."),
            ('const problem = await Problem.findById(problemId);', "Loads the target problem."),
            ('const testCases = await TestCase.find({ problemId });', "Loads hidden judge test cases."),
            ('const compilerUrl = process.env.COMPILER_SERVICE_URL || "http://localhost:7000";', "Determines compiler service base URL."),
            ('for (const tc of testCases) {', "Executes the submitted code against each test case."),
            ('const runRes = await axios.post(', "Calls compiler service `/run`."),
            ('const actual = normalize(runRes.data.output);', "Normalizes actual and expected output before comparison."),
            ('const isAccepted = passed === testCases.length;', "Computes final verdict."),
            ('problem.totalSubmissions++;', "Updates problem metrics."),
            ('const user = await User.findById(req.user._id);', "Loads user to update personal stats."),
            ('user.totalSubmissions++;', "Updates user counters and solved problems."),
            ('const submissionData = {', "Builds submission document data."),
            ('const submission = await Submission.create(submissionData);', "Persists the submission record."),
            ('return res.json({ status: submission.status });', "Returns final verdict to client."),
        ],
    },
    {
        "path": "src/Controllers/testCaseControllers.mjs",
        "title": "Test Case Controllers",
        "purpose": "Implements admin test-case creation and fetch logic.",
        "flow": "Runs after auth/admin middleware and validation, then talks directly to the `TestCase` model.",
        "points": [
            "Uses bulk insert with `insertMany`.",
            "Maps route `:id` into each new test-case document."
        ],
        "notes": ["Bulk insert is efficient, but there is no transaction around multi-document changes."],
        "anchors": [
            ('export const addTestCase = async (req, res) => {', "Starts bulk insert flow."),
            ('if (!mongoose.Types.ObjectId.isValid(id))', "Rejects malformed problem ids."),
            ('const errors = validationResult(req);', "Stops invalid input."),
            ('const { testCases } = matchedData(req, { includeOptionals: true });', "Extracts validated test-case payload."),
            ('const docs = testCases.map(tc => ({ ...tc, problemId: id }));', "Attaches parent problem id to each test case."),
            ('await TestCase.insertMany(docs);', "Bulk inserts test cases."),
            ('export const getTestCasesByProblem = async (req, res) => {', "Starts fetch-by-problem flow."),
            ('const testCases = await TestCase.find({ problemId: id });', "Reads all test cases for a problem."),
        ],
    },
    {
        "path": "src/Utils/Validating Schemas/Auth/registerValidatingSchema.mjs",
        "title": "Register Validation Schema",
        "purpose": "Validates registration input before controller logic runs.",
        "flow": "Rejects bad or duplicate registration data at route level.",
        "points": ["Uses DB-backed custom validators for unique username/email.", "Protects admin registration with `ADMIN_SECRET`."],
        "notes": ["Role is optional and defaults later in controller."],
        "anchors": [
            ('export const registerValidatingSchema = {', "Defines the registration validation schema."),
            ('userName: {', "Validates username and uniqueness."),
            ('email: {', "Validates email and uniqueness."),
            ('password: {', "Validates password length and strength."),
            ('role: {', "Validates optional role."),
            ('adminSecret: {', "Validates admin secret for admin registration."),
        ],
    },
    {
        "path": "src/Utils/Validating Schemas/Auth/loginValidatingSchema.mjs",
        "title": "Login Validation Schema",
        "purpose": "Validates login credentials and preloads the matched user.",
        "flow": "Runs before login controller and stores the found user on `req._foundUser`.",
        "points": ["Validator performs both account lookup and password comparison.", "Optional role check prevents logging in under the wrong declared role."],
        "notes": ["This pattern keeps the controller small by pushing credential checks into validation."],
        "anchors": [
            ('export const loginValidatingSchema = {', "Defines the login validation schema."),
            ('email: {', "Validates email and loads the matching user."),
            ('password: {', "Compares submitted password with stored hash."),
            ('role: {', "Validates optional role against stored account role."),
            ('export default loginValidatingSchema', "Exports the schema."),
        ],
    },
    {
        "path": "src/Utils/Validating Schemas/problemValidatingSchema.mjs",
        "title": "Problem Validation Schema",
        "purpose": "Validates problem-creation payloads.",
        "flow": "Ensures only well-formed problem data reaches `addProblem` controller.",
        "points": ["Uses nested validation for examples array.", "Checks uniqueness of both problem number and title."],
        "notes": ["Converts `problemNumber` to int before controller use."],
        "anchors": [
            ('export const problemValidatingSchema = {', "Defines the problem validation schema."),
            ('problemNumber: {', "Validates integer problem number and uniqueness."),
            ('title: {', "Validates title and uniqueness."),
            ('statement: {', "Validates problem statement."),
            ('examples: {', "Validates examples array."),
            ('"examples.*.input": {', "Validates each example input."),
            ('"examples.*.output": {', "Validates each example output."),
            ('constraints: {', "Validates constraints."),
        ],
    },
    {
        "path": "src/Utils/Validating Schemas/profileValidatingSchema.mjs",
        "title": "Profile Validation Schema",
        "purpose": "Validates incoming profile updates.",
        "flow": "Runs before profile update controller.",
        "points": ["All fields are optional to support partial updates.", "Profile picture is validated as a URL."],
        "notes": ["`firstName` and `lastName` are validated but never stored."],
        "anchors": [
            ('export const profileValidatingSchema = {', "Defines the profile validation schema."),
            ('firstName: {', "Validates optional first name."),
            ('lastName: {', "Validates optional last name."),
            ('profilePicture: {', "Validates profile image URL."),
            ('bio: {', "Validates bio."),
            ('location: {', "Validates location."),
            ('institution: {', "Validates institution."),
        ],
    },
    {
        "path": "src/Utils/Validating Schemas/submissionValidatingSchema.mjs",
        "title": "Submission Validation Schema",
        "purpose": "Validates code and language for code-submission requests.",
        "flow": "Runs before expensive judge logic starts.",
        "points": ["Early validation matters because submit flow is costly.", "Problem id comes from URL, so body only needs code and language."],
        "notes": ["Language is not restricted to a fixed allow-list here."],
        "anchors": [
            ('export const submissionValidatingSchema = {', "Defines submission validation."),
            ('code: {', "Validates code presence and type."),
            ('language: {', "Validates language presence and type."),
        ],
    },
    {
        "path": "src/Utils/Validating Schemas/testCaseValidatingSchema.mjs",
        "title": "Test Case Validation Schema",
        "purpose": "Validates admin test-case payloads before bulk insert.",
        "flow": "Runs before `addTestCase` controller.",
        "points": ["Validates both top-level array and nested test-case fields.", "Keeps controller logic simple by guaranteeing payload shape."],
        "notes": [],
        "anchors": [
            ('export const testCaseValidatingSchema = {', "Defines test-case validation."),
            ('"testCases": {', "Validates the array container."),
            ('"testCases.*.input": {', "Validates each input field."),
            ('"testCases.*.output": {', "Validates each output field."),
        ],
    },
])

COMPILER_GUIDES.extend([
    {
        "path": "Index.js",
        "title": "Compiler Service Entrypoint",
        "purpose": "Receives code-execution requests from backend and dispatches execution by language.",
        "flow": "Backend submission controller calls `POST /run` here for every hidden test case.",
        "anchors": [
            ("const express = require('express');", "Bootstraps the compiler service and loads helpers."),
            ('app.get("/", (req, res) => {', "Defines a health endpoint."),
            ('app.post("/run", async (req, res) => {', "Defines the execution endpoint used by the backend."),
            ('const filePath = await generateFile(language, code);', "Writes source code to disk."),
            ('const inputPath = await generateInputFile(input);', "Writes stdin data to disk."),
            ('if (language === "cpp") {', "Branches to the appropriate runtime handler."),
            ('res.json({', "Returns output to the backend."),
            ('app.post("/ai-review", async (req, res) => {', "Defines an AI review endpoint separate from judge execution."),
        ],
    },
    {
        "path": "generateFile.js",
        "title": "Compiler Source File Generator",
        "purpose": "Creates temporary source files for the compiler service.",
        "flow": "Called by compiler service before code execution.",
        "anchors": [
            ("const fs = require('fs');", "Imports filesystem, path, and uuid helpers."),
            ("const dirCodes = path.join(__dirname, 'codes');", "Sets the source-code folder."),
            ("if (!fs.existsSync(dirCodes)) {", "Creates folder if needed."),
            ('const generateFile = (format, content) => {', "Defines file-creation helper."),
            ('const jobID = uuid();', "Creates a unique job id."),
            ('fs.writeFileSync(filePath, content);', "Writes code file to disk."),
        ],
    },
    {
        "path": "generateInputFile.js",
        "title": "Compiler Input File Generator",
        "purpose": "Creates temporary input files for program stdin.",
        "flow": "Called by compiler service before execution.",
        "anchors": [
            ("const fs = require('fs');", "Imports filesystem helpers."),
            ("const dirInputs = path.join(__dirname, 'inputs');", "Sets the input folder."),
            ("if (!fs.existsSync(dirInputs)) {", "Creates folder if needed."),
            ('const generateInputFile = async (input) => {', "Defines helper for writing input files."),
            ('await fs.writeFileSync(input_filePath, input);', "Writes input contents to disk."),
        ],
    },
])

GUIDES.extend([
    {
        "path": "src/Models/User.mjs",
        "title": "User Model",
        "purpose": "Stores login identity, role, profile information, and progress metrics.",
        "flow": "Used by auth, profile, and submission logic across the whole backend.",
        "points": [
            "Combines auth fields and gamification/profile fields in one document.",
            "Unique `userName` and `email` enforce account uniqueness."
        ],
        "notes": ["The schema has `userName`, not `name`, which matters because JWT helper references `user.name`."],
        "anchors": [
            ('const userSchema = new mongoose.Schema({', "Starts the user schema definition."),
            ('userName: {', "Stores the public username/handle."),
            ('email: {', "Stores the login email with uniqueness."),
            ('password: {', "Stores hashed password data."),
            ('role: {', "Stores `user` or `admin`."),
            ('profilePicture: {', "Begins optional profile fields."),
            ('problemsSolved: {', "Stores solved problem numbers and aggregate stats."),
            ("userSchema.pre('save', function (next) {", "Normalizes role to lowercase before save."),
            ('export const User = mongoose.models.User || mongoose.model("User", userSchema);', "Exports the reusable model."),
        ],
    },
    {
        "path": "src/Models/Problem.mjs",
        "title": "Problem Model",
        "purpose": "Stores coding problem content, examples, constraints, creator, and acceptance stats.",
        "flow": "Used by problem controllers and updated again by the submission controller.",
        "points": [
            "Examples are embedded subdocuments because they belong tightly to a problem.",
            "Acceptance rate is a derived field stored for read efficiency."
        ],
        "notes": ["Deleting a problem requires manual cleanup of test cases and submissions."],
        "anchors": [
            ('const exampleSchema = new mongoose.Schema({', "Defines the embedded example structure."),
            ('const problemSchema = new mongoose.Schema({', "Defines the problem document."),
            ('problemNumber: {', "Stores a unique numeric identifier."),
            ('title: {', "Stores a unique problem title."),
            ('statement: {', "Stores the problem statement."),
            ('examples: {', "Stores example input/output pairs."),
            ('constraints: {', "Stores textual constraints."),
            ('totalSubmissions: {', "Begins stored statistics fields."),
            ('created_by: {', "Links the problem to the admin who created it."),
            ('export const Problem = mongoose.models.Problem || mongoose.model("Problem", problemSchema);', "Exports the Problem model."),
        ],
    },
    {
        "path": "src/Models/Testcase.mjs",
        "title": "Test Case Model",
        "purpose": "Stores hidden input/output data used by the judge.",
        "flow": "Loaded by submission controller and managed through admin-only test-case routes.",
        "points": [
            "Hidden judge data is separated from public problem content.",
            "Each test case points to its parent problem through `problemId`."
        ],
        "notes": ["An index on `problemId` would help repeated lookup during judging."],
        "anchors": [
            ('const testCaseSchema = new mongoose.Schema({', "Defines the test-case schema."),
            ('problemId: {', "Links the test case to a problem."),
            ('input: {', "Stores stdin-style input."),
            ('output: {', "Stores expected output."),
            ('export const TestCase = mongoose.models.TestCase || mongoose.model("TestCase", testCaseSchema);', "Exports the TestCase model."),
        ],
    },
    {
        "path": "src/Models/Submission.mjs",
        "title": "Submission Model",
        "purpose": "Stores each user submission with code, verdict, related ids, and execution metadata.",
        "flow": "Written by `submitSolution` and read by submission history endpoints.",
        "points": [
            "Submission documents create an audit trail of every attempt.",
            "Stored `problemNumber` avoids needing a problem join for simple history views."
        ],
        "notes": [
            "`submissionId` based on `Date.now()` is not a strong unique identifier under concurrency.",
            "Indexes on user/problem lookup paths would help performance."
        ],
        "anchors": [
            ('const submissionSchema = new mongoose.Schema({', "Defines the submission document."),
            ('submissionId: {', "Adds an app-level timestamp-based id."),
            ('userId: {', "Links submission to user."),
            ('problemId: {', "Links submission to problem."),
            ('problemNumber: {', "Stores a copied problem number."),
            ('language: {', "Stores submission language."),
            ('code: {', "Stores submitted source code."),
            ('input: {', "Stores judge input snapshot."),
            ('status: {', "Stores result verdict."),
            ('createdAt: {', "Stores creation timestamp."),
            ('export const Submission = mongoose.models.Submission || mongoose.model("Submission", submissionSchema);', "Exports the model."),
        ],
    },
    {
        "path": "src/Routes/authRoutes.mjs",
        "title": "Auth Routes",
        "purpose": "Maps auth URLs to auth validation and controller functions.",
        "flow": "Incoming auth requests are dispatched here after global middleware completes in `Index.mjs`.",
        "points": [
            "Validation is attached at route level with `checkSchema(...)`.",
            "Current-user lookup is protected by `requireAuth`."
        ],
        "notes": ["Router uses full paths rather than a mounted `/api/auth` prefix."],
        "anchors": [
            ('const router = Router();', "Creates the router instance."),
            ("router.post('/api/auth/register'", "Register route with validation."),
            ("router.post('/api/auth/login'", "Login route with validation."),
            ("router.get('/api/auth/me'", "Authenticated current-user route."),
            ("router.post('/api/auth/logout'", "Logout route."),
            ('export default router;', "Exports the router."),
        ],
    },
    {
        "path": "src/Routes/problemRoutes.mjs",
        "title": "Problem Routes",
        "purpose": "Maps problem listing, read, create, and delete endpoints.",
        "flow": "Problem routes call public controllers directly or run auth/admin/validation middleware first.",
        "points": [
            "Middleware order is `requireAuth` then `isAdmin` for sensitive routes.",
            "Route design mixes plural collection path and singular resource path."
        ],
        "notes": ["Single problem read is protected, while problem list is public."],
        "anchors": [
            ('const router = Router();', "Creates the problem router."),
            ("router.get('/api/problems'", "Public list route."),
            ("router.get('/api/problem/:Id'", "Authenticated single-problem route."),
            ("router.post('/api/problem'", "Admin-only create route."),
            ("router.delete('/api/problem/:Id'", "Admin-only delete route."),
            ('export default router;', "Exports the router."),
        ],
    },
    {
        "path": "src/Routes/testCaseRoutes.mjs",
        "title": "Test Case Routes",
        "purpose": "Maps admin-only test-case creation and read endpoints.",
        "flow": "Every request on this router passes auth and admin authorization before touching hidden judge data.",
        "points": [
            "Keeping hidden test cases admin-only protects judge integrity.",
            "Bulk insert route validates the request payload before controller execution."
        ],
        "notes": ["Uses `:id` param casing unlike problem routes' `:Id`."],
        "anchors": [
            ('const router = Router();', "Creates the test-case router."),
            ('router.post("/api/testcases/:id"', "Admin-only create route."),
            ('router.get("/api/testcases/:id"', "Admin-only read route."),
            ('export default router;', "Exports the router."),
        ],
    },
    {
        "path": "src/Routes/profileRoutes.mjs",
        "title": "Profile Routes",
        "purpose": "Maps leaderboard, public profile, self-profile, and profile update endpoints.",
        "flow": "Public profile routes bypass auth while self-profile routes require it.",
        "points": [
            "Leaderboard is intentionally public.",
            "Profile update uses validation before controller logic."
        ],
        "notes": ["`/api/profile` is the current-user resource; `/api/user/:id` is a public user resource."],
        "anchors": [
            ('const router = Router();', "Creates the profile router."),
            ('router.get("/api/leaderboard"', "Public leaderboard route."),
            ('router.get("/api/user/:id"', "Public user profile route."),
            ('router.get("/api/profile"', "Authenticated own-profile route."),
            ('router.put("/api/profile"', "Authenticated profile update route."),
            ('export default router;', "Exports the router."),
        ],
    },
    {
        "path": "src/Routes/submissionRoutes.mjs",
        "title": "Submission Routes",
        "purpose": "Maps submission history and code-judging endpoints.",
        "flow": "All routes require auth; the submit route also validates request body and triggers the deepest controller flow in the backend.",
        "points": [
            "These routes are performance-sensitive because submit path touches DB and external compiler service.",
            "History endpoints scope results to the logged-in user."
        ],
        "notes": ["Submit route is where most interview discussion about request flow should focus."],
        "anchors": [
            ('const router = Router();', "Creates the submission router."),
            ('router.get("/api/submissions"', "All-submissions history route."),
            ('router.get("/api/problem/:id/submissions"', "Per-problem history route."),
            ('router.post("/api/problem/:id/submit"', "Judging route."),
            ('export default router;', "Exports the router."),
        ],
    },
])


def add_front(doc):
    title(doc, "Code Bharat Backend Interview Guide")
    title(doc, f"Generated from this repository on {datetime.now().strftime('%Y-%m-%d %H:%M')}", 11, True)
    doc.add_paragraph("Goal: explain the real backend codebase deeply enough for interviews, debugging, and confident system design discussion.")
    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph("This backend is a Node.js + Express + MongoDB API for a coding platform. It manages authentication, problems, hidden test cases, profile and leaderboard features, and solution judging through a separate compiler service.")
    doc.add_heading("2. High-Level Flow", level=1)
    code(doc, "\n".join([
        "Client",
        "  -> Express app (Index.mjs)",
        "    -> Global middleware (CORS, JSON parser, URL-encoded parser, cookie parser)",
        "      -> Router",
        "        -> Route middleware (validation, requireAuth, isAdmin)",
        "          -> Controller",
        "            -> Mongoose model / MongoDB",
        "            -> External compiler service (submissions only)",
        "          -> JSON response",
    ]))
    doc.add_heading("3. Request Travel Example: Submit Solution", level=1)
    nums(doc, [
        "Frontend sends `POST /api/problem/:id/submit` with code and language.",
        "Express runs CORS and parsers, then reads cookies.",
        "`requireAuth` verifies the JWT cookie and loads the user.",
        "`checkSchema(submissionValidatingSchema)` validates code and language.",
        "`submitSolution` loads the problem and hidden test cases from MongoDB.",
        "For each hidden test case, backend calls compiler service `/run`.",
        "Backend compares normalized output with expected output.",
        "Backend updates problem stats, user stats, and writes a submission record.",
        "Backend returns the final verdict JSON.",
    ])
    doc.add_heading("4. Folder Structure", level=1)
    code(doc, tree())
    doc.add_heading("5. API Surface", level=1)
    table(doc, ["Method", "Path", "Access", "Chain"], ROUTES)
    doc.add_heading("6. Environment Variables", level=1)
    table(doc, ["Variable", "Use"], ENV_ROWS)
    doc.add_heading("7. Key Interview Risks / Observations", level=1)
    bullets(doc, [
        "No centralized error middleware.",
        "Submission stats and submission creation are not transactional.",
        "`jwt.mjs` has `dotenv.config;` instead of `dotenv.config();`.",
        "JWT payload uses `user.name`, but schema field is `userName`.",
        "Profile validator includes `firstName` and `lastName`, but model/controller do not persist them.",
        "Submission history endpoints do not use `try/catch`.",
        "There are no explicit performance indexes for heavy lookup paths like submissions and test cases.",
    ])


def add_env_page(doc):
    env = BACKEND / ".env"
    if not env.exists():
        return
    redacted = []
    for line in read(env).splitlines():
        if not line.strip() or line.lstrip().startswith("#"):
            redacted.append(line)
        else:
            redacted.append(line.split("=", 1)[0] + "=***REDACTED***")
    doc.add_page_break()
    doc.add_heading("Redacted `.env`", level=1)
    code(doc, numbered("\n".join(redacted)))


def add_file(doc, entry):
    path = BACKEND / entry["path"]
    text = read(path)
    doc.add_page_break()
    doc.add_heading(entry["title"], level=1)
    doc.add_paragraph(f"File: Backend/{entry['path']}")
    doc.add_paragraph(f"Purpose: {entry['purpose']}")
    doc.add_paragraph(f"How it is used in request flow: {entry['flow']}")
    doc.add_paragraph("What to say in interview:")
    bullets(doc, entry["points"])
    if entry.get("notes"):
        doc.add_paragraph("Implementation notes:")
        bullets(doc, entry["notes"])
    doc.add_heading("Code", level=2)
    code(doc, numbered(text))
    doc.add_heading("Walkthrough", level=2)
    walkthrough(doc, text, entry["anchors"])


def add_compiler(doc, entry):
    path = COMPILER / entry["path"]
    text = read(path)
    doc.add_heading(entry["title"], level=2)
    doc.add_paragraph(f"File: Compiler/{entry['path']}")
    doc.add_paragraph(f"Purpose: {entry['purpose']}")
    doc.add_paragraph(f"Backend dependency: {entry['flow']}")
    code(doc, numbered(text))
    walkthrough(doc, text, entry["anchors"])


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style(doc)
    add_front(doc)
    add_env_page(doc)
    for entry in GUIDES:
        add_file(doc, entry)
    doc.add_page_break()
    doc.add_heading("Compiler Service Appendix", level=1)
    doc.add_paragraph("The backend submission flow depends on a separate compiler service. Understanding the contract helps you explain how code travels beyond the main API.")
    for entry in COMPILER_GUIDES:
        add_compiler(doc, entry)
    doc.add_page_break()
    doc.add_heading("Best Practices To Mention In Interviews", level=1)
    bullets(doc, [
        "Explain the code as layered flow: entrypoint -> middleware -> routes -> controllers -> models -> database.",
        "Mention cookie-based JWT auth, not localStorage token auth.",
        "Mention that hidden judge test cases are separated from visible problem content.",
        "Mention that submission orchestration is the most complex and performance-sensitive path.",
        "Mention improvements: transactions, indexes, global error middleware, rate limiting, logging, and tests.",
    ])
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

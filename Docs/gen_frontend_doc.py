from pathlib import Path
from datetime import datetime
from docx import Document
from gen_doc import style, title, bullets, nums, code, numbered, table, walkthrough, read

ROOT = Path(__file__).resolve().parent
FRONTEND = ROOT / "Frontend"
OUT = ROOT / "output" / "doc" / "Code-Bharat-Frontend-Interview-Guide.docx"


def tree():
    files = sorted(
        p.relative_to(FRONTEND).as_posix()
        for p in FRONTEND.rglob("*")
        if p.is_file()
        and "node_modules" not in p.parts
        and "dist" not in p.parts
        and "assets" not in p.parts
        and p.name != "package-lock.json"
    )
    out = ["Frontend/"]
    for f in files:
        out.append(f"{'  ' * f.count('/')} - {f.split('/')[-1]}")
    return "\n".join(out)


GUIDES = []

GUIDES.extend([
    {
        "path": "package.json",
        "title": "Frontend Package Manifest",
        "purpose": "Defines the frontend stack, scripts, runtime mode, and key dependencies.",
        "flow": "This file decides how the frontend boots, builds, lints, and which libraries are available to the app.",
        "points": [
            "Vite is the build/dev tool, React 19 is the UI runtime, and Tailwind is integrated through the Vite plugin.",
            "React Router, Redux Toolkit, axios, Monaco, React Split, React Markdown, and Framer Motion are core libraries."
        ],
        "notes": ["This is a Vite ESM app because `type` is `module`."],
        "anchors": [
            ('"type": "module"', "Marks the project as ES module based."),
            ('"scripts": {', "Defines development, build, lint, and preview commands."),
            ('"dependencies": {', "Lists runtime libraries used by the UI."),
            ('"devDependencies": {', "Lists development tooling like Vite and ESLint."),
        ],
    },
    {
        "path": "vite.config.js",
        "title": "Vite Config",
        "purpose": "Registers build-time plugins for React and Tailwind CSS.",
        "flow": "Used during development and build rather than browser runtime execution.",
        "points": [
            "React plugin enables JSX and React-specific Vite behavior.",
            "Tailwind plugin integrates utility-first styling into the build."
        ],
        "notes": ["Config is intentionally minimal."],
        "anchors": [
            ("import { defineConfig } from 'vite'", "Imports Vite config builder and plugins."),
            ('export default defineConfig({', "Starts the Vite config object."),
            ('plugins: [react(),', "Registers React and Tailwind plugins."),
        ],
    },
    {
        "path": "eslint.config.js",
        "title": "ESLint Config",
        "purpose": "Defines lint rules for JavaScript and JSX files.",
        "flow": "Affects development quality checks rather than runtime logic.",
        "points": [
            "Uses flat ESLint config format.",
            "Includes React hooks and React Refresh recommendations."
        ],
        "notes": ["Built `dist` output is globally ignored."],
        "anchors": [
            ("import js from '@eslint/js'", "Imports ESLint presets and helper packages."),
            ('export default defineConfig([', "Defines the flat ESLint config array."),
            ("globalIgnores(['dist'])", "Ignores generated build output."),
            ("files: ['**/*.{js,jsx}']", "Applies rules to JS and JSX files."),
            ("rules: {", "Defines project-specific rule overrides."),
        ],
    },
    {
        "path": "index.html",
        "title": "HTML Shell",
        "purpose": "Defines the single HTML page and the root mount target for React.",
        "flow": "Browser loads this file first and then runs `src/main.jsx`.",
        "points": [
            "This is a standard SPA shell with one root div.",
            "Favicon is loaded from a frontend asset."
        ],
        "notes": [],
        "anchors": [
            ('<!doctype html>', "Declares the HTML document."),
            ('<link rel="icon"', "Sets favicon metadata."),
            ('<div id="root"></div>', "Creates the DOM mount point."),
            ('<script type="module" src="/src/main.jsx"></script>', "Starts the React app from `main.jsx`."),
        ],
    },
    {
        "path": "src/main.jsx",
        "title": "React Entrypoint",
        "purpose": "Mounts the app into the DOM and wraps it in Redux Provider.",
        "flow": "First JavaScript module executed by the browser after `index.html`.",
        "points": [
            "Redux Provider makes auth state available to the whole app.",
            "StrictMode helps surface side-effect issues during development."
        ],
        "notes": ["Global styles come from `index.css`."],
        "anchors": [
            ("import { StrictMode } from 'react';", "Imports React root helpers, Provider, store, App, and global CSS."),
            ("createRoot(document.getElementById('root')).render(", "Mounts the React tree into the DOM."),
            ('<StrictMode>', "Wraps the app in React StrictMode."),
            ('<Provider store={store}>', "Injects the Redux store."),
            ('<App />', "Renders the router app."),
        ],
    },
    {
        "path": "src/App.jsx",
        "title": "Router Composition",
        "purpose": "Defines the route tree, layouts, loaders, and protected admin section.",
        "flow": "Decides which layout and page render for every URL in the application.",
        "points": [
            "Public, problem workspace, auth, and admin flows are separated with different layouts.",
            "Admin routes are protected with `RequireAdmin` at route level.",
            "Problem pages use React Router loaders."
        ],
        "notes": ["There is no catch-all route, so `NotFoundPage` is currently unused."],
        "anchors": [
            ('function App() {', "Defines the top-level app component."),
            ('const router = createBrowserRouter(', "Creates the browser router."),
            ('<Route path="/" element={<RootLayout />}>', "Defines public root layout routes."),
            ('<Route element={<ProblemLayout />}>', "Defines problem workspace routes."),
            ('<Route element={<LoginLayout />}>', "Defines user auth routes."),
            ('<Route element={<AdminAuthLayout />}>', "Defines admin auth routes."),
            ('<Route path="/admin" element={<RequireAdmin />}>', "Defines protected admin routes."),
            ('return <RouterProvider router={router} />;', "Renders router provider."),
        ],
    },
    {
        "path": "src/App.css",
        "title": "App CSS",
        "purpose": "Placeholder stylesheet with no active content.",
        "flow": "Does not currently affect rendering because it is empty and unused.",
        "points": ["This file is effectively unused in the current frontend."],
        "notes": ["Shared styling actually lives in `src/index.css` and Tailwind classes."],
        "anchors": [],
    },
    {
        "path": "src/index.css",
        "title": "Global CSS",
        "purpose": "Loads Tailwind and defines split-pane gutter styling.",
        "flow": "Imported once in `main.jsx` and applied globally.",
        "points": [
            "Tailwind is enabled here with `@import`.",
            "Custom gutter styles support `react-split` in the coding workspace."
        ],
        "notes": [],
        "anchors": [
            ('@import "tailwindcss";', "Loads Tailwind CSS."),
            ('.gutter {', "Defines base split-gutter styling."),
            ('.gutter.gutter-horizontal {', "Styles horizontal resize gutter."),
            ('.gutter.gutter-vertical {', "Styles vertical resize gutter."),
        ],
    },
    {
        "path": "src/config/urls.js",
        "title": "URL Config",
        "purpose": "Computes backend and compiler-service base URLs from Vite environment variables.",
        "flow": "Imported before axios calls are made from pages and API helpers.",
        "points": [
            "Frontend talks to two services: backend API and compiler service.",
            "Development and production URLs are resolved separately."
        ],
        "notes": [],
        "anchors": [
            ('const localBackendUrl =', "Defines development backend default or override."),
            ('const localCompilerUrl =', "Defines development compiler default or override."),
            ('export const BACKEND_URL = import.meta.env.DEV', "Chooses backend URL by environment."),
            ('export const COMPILER_URL = import.meta.env.DEV', "Chooses compiler URL by environment."),
        ],
    },
    {
        "path": "src/Hooks/useAuth.js",
        "title": "useAuth Hook",
        "purpose": "Provides a small hook for fetching current auth state from the backend.",
        "flow": "Could be used by any component needing auth, although Redux auth is more common in this repo.",
        "points": [
            "Uses local state for `user` and `loading`.",
            "Relies on backend cookie auth with `withCredentials: true`."
        ],
        "notes": ["Appears unused in the current codebase."],
        "anchors": [
            ('export default function useAuth() {', "Defines the auth hook."),
            ('const [user, setUser] = useState(null);', "Stores local auth user."),
            ('useEffect(() => {', "Fetches current user on mount."),
            ('const res = await axios.get(', "Calls `/api/auth/me`."),
            ('return { user, loading };', "Returns hook state."),
        ],
    },
    {
        "path": "src/Store/App/store.js",
        "title": "Redux Store",
        "purpose": "Creates the global Redux store.",
        "flow": "Provided at app boot and read by navbar, admin guard, and other UI pieces.",
        "points": [
            "Redux scope is intentionally narrow: only auth is global.",
            "Auth state is exposed under `state.auth`."
        ],
        "notes": [],
        "anchors": [
            ('export const store = configureStore({', "Creates the Redux store."),
            ('reducer: {', "Registers reducers."),
            ('auth: authReducer,', "Mounts auth reducer."),
        ],
    },
    {
        "path": "src/Store/Features/Auth/authAPI.js",
        "title": "Auth API Helpers",
        "purpose": "Wraps auth-related axios requests used by thunks.",
        "flow": "Thunk layer calls these helpers when UI needs current-user or logout operations.",
        "points": [
            "Separates raw HTTP calls from Redux thunk declarations.",
            "Both calls send cookies with `withCredentials: true`."
        ],
        "notes": [],
        "anchors": [
            ('export const getUser = async () => {', "Defines current-user request helper."),
            ('const res = await axios.get(`${BACKEND_URL}/api/auth/me`', "Calls backend current-user endpoint."),
            ('export const logout = async () => {', "Defines logout request helper."),
            ('await axios.post(', "Calls backend logout endpoint."),
        ],
    },
    {
        "path": "src/Store/Features/Auth/authSlice.js",
        "title": "Auth Slice",
        "purpose": "Stores authenticated user data and loading state in Redux.",
        "flow": "Global auth state drives navbar rendering and admin route protection.",
        "points": [
            "Async state is handled through `extraReducers`.",
            "Rejected auth fetch clears the user."
        ],
        "notes": [],
        "anchors": [
            ('const initialState = {', "Defines default auth state."),
            ('const authSlice = createSlice({', "Creates the auth slice."),
            ('.addCase(fetchUser.pending, (state) => {', "Marks loading while auth check is in progress."),
            ('.addCase(fetchUser.fulfilled, (state, action) => {', "Stores current user on success."),
            ('.addCase(fetchUser.rejected, (state) => {', "Clears user on failed auth check."),
            ('.addCase(logoutThunk.fulfilled, (state) => {', "Clears user on logout."),
        ],
    },
    {
        "path": "src/Store/Features/Auth/authThunks.js",
        "title": "Auth Thunks",
        "purpose": "Defines async thunks for fetching current user and logging out.",
        "flow": "Components dispatch these thunks to keep auth state synchronized with backend.",
        "points": [
            "`fetchUser` hydrates auth state for navbar and guards.",
            "`logoutThunk` keeps Redux state aligned with server logout."
        ],
        "notes": [],
        "anchors": [
            ('import { createAsyncThunk } from "@reduxjs/toolkit";', "Imports thunk helper."),
            ("export const fetchUser = createAsyncThunk('/auth/fetchUser', getUser);", "Defines current-user thunk."),
            ("export const logoutThunk = createAsyncThunk('/auth/logout', logout);", "Defines logout thunk."),
        ],
    },
])

GUIDES.extend([
    {
        "path": "src/Pages/Admin/AdminDashboardPage.jsx",
        "title": "Admin Dashboard Page",
        "purpose": "Landing page for the protected admin workspace.",
        "flow": "Rendered at `/admin` after `RequireAdmin` and inside `AdminLayout`.",
        "points": ["Acts as navigation hub for admin problem management features."],
        "notes": [],
        "anchors": [
            ('export default function AdminDashboardPage() {', "Defines admin dashboard page."),
            ('<Link', "Renders admin action cards linking to problem management and creation."),
        ],
    },
    {
        "path": "src/Pages/Admin/AdminProblemsPage.jsx",
        "title": "Admin Problems Page",
        "purpose": "Lets admins view the problem bank and delete problems.",
        "flow": "Rendered at `/admin/problems` inside protected admin workspace.",
        "points": [
            "Fetches problems from backend and displays admin controls.",
            "Delete action calls backend admin delete endpoint and updates local UI state."
        ],
        "notes": ["Uses browser `window.confirm` for delete confirmation."],
        "anchors": [
            ('export default function AdminProblemsPage() {', "Defines admin problem-management page."),
            ('const [problems, setProblems] = useState([]);', "Stores fetched problems."),
            ('const fetchProblems = async () => {', "Defines problem fetch helper."),
            ('useEffect(() => {', "Fetches problems on mount."),
            ('const handleDelete = async (id) => {', "Deletes a problem after confirmation."),
            ('return (', "Renders management table UI."),
        ],
    },
    {
        "path": "src/Pages/Admin/AdminAddProblemPage.jsx",
        "title": "Admin Add Problem Page",
        "purpose": "Implements the full admin workflow for creating a problem, visible examples, and hidden test cases.",
        "flow": "Rendered at `/admin/problems/new`; first posts problem to backend, then posts hidden test cases to a second endpoint.",
        "points": [
            "This page shows how the frontend coordinates a two-step admin workflow with dynamic example/test-case rows.",
            "Visible examples and hidden test cases are kept in separate local arrays."
        ],
        "notes": [
            "Because problem creation and test-case creation are separate API calls, a failure in the second call can leave a partially created problem.",
            "The page uses imperative `alert` messaging for success and failure feedback."
        ],
        "anchors": [
            ('function AdminAddProblemPage() {', "Defines admin problem-creation page."),
            ('const [form, setForm] = useState({', "Stores top-level problem fields."),
            ('const [examples, setExamples] = useState([{ input: "", output: "" }]);', "Stores visible example rows."),
            ('const [testCases, setTestCases] = useState([{ input: "", output: "" }]);', "Stores hidden test-case rows."),
            ('const handleChange = (e) =>', "Updates top-level form state."),
            ('const updateArray =', "Builds reusable row update helper."),
            ('const addRow = (setter) => () =>', "Builds reusable row append helper."),
            ('const formValid =', "Calculates overall form validity."),
            ('const handleSubmit = async (e) => {', "Runs the two-step submit workflow."),
            ('const problemResponse = await axios.post(', "Creates the problem first."),
            ('await axios.post(', "Creates hidden test cases second."),
            ('return (', "Renders admin form UI."),
            ('const Section = ({ title, description, rows, onAdd, onChange }) => (', "Defines reusable array-section UI block."),
        ],
    },
    {
        "path": "src/Pages/Admin/AdminLoginPage.jsx",
        "title": "Admin Login Page",
        "purpose": "Implements admin login UI and submits admin-role credentials to the shared backend auth endpoint.",
        "flow": "Rendered at `/admin/login` inside `AdminAuthLayout`.",
        "points": [
            "Uses same backend login endpoint as user login but sends `role: 'admin'`.",
            "Redirects into protected admin workspace on success."
        ],
        "notes": [],
        "anchors": [
            ('export default function AdminLoginPage() {', "Defines admin login page."),
            ('const [formFields, setFormFields] = useState({ email: "", password: "" });', "Stores login form state."),
            ('useEffect(() => {', "Runs submit-side effect when submit flag changes."),
            ('await axios.post(', "Calls backend login endpoint with admin role."),
            ('return (', "Renders admin login UI."),
        ],
    },
    {
        "path": "src/Pages/Admin/AdminRegisterPage.jsx",
        "title": "Admin Register Page",
        "purpose": "Implements admin registration UI with admin secret and stronger password requirements.",
        "flow": "Rendered at `/admin/register` inside `AdminAuthLayout`.",
        "points": [
            "Uses same backend register endpoint as user register but sends `role: 'admin'` plus `adminSecret`.",
            "Separates admin registration flow from public user registration UI."
        ],
        "notes": ["This page reinforces the backend design where admin accounts require a secret."],
        "anchors": [
            ('export default function AdminRegisterPage() {', "Defines admin register page."),
            ('const [formFields, setFormFields] = useState({', "Stores admin registration form state."),
            ('const passwordRequirements =', "Defines visible password rules."),
            ('const isStrongPassword = (password) =>', "Defines client-side password rule."),
            ('useEffect(() => {', "Runs submit-side effect when submit flag changes."),
            ('await axios.post(', "Calls backend register endpoint with admin payload."),
            ('return (', "Renders admin registration UI."),
        ],
    },
])

GUIDES.extend([
    {
        "path": "src/Pages/HomePage.jsx",
        "title": "Home Page",
        "purpose": "Landing page that sets the visual tone and funnels users into main app areas.",
        "flow": "Rendered at `/` inside `RootLayout`.",
        "points": [
            "Uses Framer Motion for gentle hero animation and staged reveals.",
            "Reads Redux auth state to switch CTA between profile and registration."
        ],
        "notes": ["Design uses the mesh background plus overlay gradients as the main visual identity."],
        "anchors": [
            ('const links = [', "Defines hero quick-link cards."),
            ('export default function HomePage() {', "Defines the landing page component."),
            ('const { user } = useSelector((state) => state.auth);', "Reads auth state to personalize CTA."),
            ('<motion.img', "Animates the mesh background."),
            ('<motion.h1', "Renders the animated hero headline."),
            ('{user ? (', "Switches CTA based on login state."),
            ('{links.map((item) => (', "Renders landing-page navigation cards."),
        ],
    },
    {
        "path": "src/Pages/ExplorePage.jsx",
        "title": "Explore Page",
        "purpose": "Searchable discovery page for browsing problems by title, number, acceptance band, and sorting preference.",
        "flow": "Rendered at `/explore` inside `RootLayout` with `MeshPageShell`.",
        "points": [
            "Fetches the problem list from backend on mount.",
            "Uses `useMemo` for filtered and sorted derived data.",
            "Design emphasizes discovery through stats cards, filter pills, and spotlight links."
        ],
        "notes": [],
        "anchors": [
            ('const acceptanceFilters = [', "Defines the filter pill options."),
            ('export default function ExplorePage() {', "Defines the explore page."),
            ('const [problems, setProblems] = useState([]);', "Stores fetched problems."),
            ('useEffect(() => {', "Fetches problem data from backend."),
            ('const filteredProblems = useMemo(() => {', "Computes search/filter/sort result list."),
            ('const averageAcceptance = problems.length', "Computes summary metrics."),
            ('return (', "Renders the page layout."),
            ('function StatCard({ label, value }) {', "Defines reusable inline stat card."),
            ('function Panel({ title, children }) {', "Defines reusable side panel."),
            ('function BandRow({ label, count, tone }) {', "Defines band summary row."),
        ],
    },
    {
        "path": "src/Pages/ContestsPage.jsx",
        "title": "Contests Page",
        "purpose": "Builds practice contest packs from the live problem bank instead of using a separate contest API.",
        "flow": "Rendered at `/contests` inside `RootLayout` with `MeshPageShell`.",
        "points": [
            "Fetches problems once and derives different contest sets with `useMemo`.",
            "Demonstrates a frontend-only composition feature built on top of existing backend problem data."
        ],
        "notes": ["Contest packs are heuristic groupings by acceptance rate, not backend-defined contests."],
        "anchors": [
            ('export default function ContestsPage() {', "Defines the contests page."),
            ('const [problems, setProblems] = useState([]);', "Stores fetched problems."),
            ('useEffect(() => {', "Fetches problem bank from backend."),
            ('const contestCards = useMemo(() => {', "Builds derived contest packs from problem data."),
            ('return (', "Renders contest page layout."),
            ('function MiniMetric({ label, value }) {', "Defines inline metric card."),
        ],
    },
    {
        "path": "src/Pages/LeaderboardPage.jsx",
        "title": "Leaderboard Page",
        "purpose": "Shows ranked user standings with podium cards and a full table.",
        "flow": "Rendered at `/leaderboard` inside `RootLayout` with `MeshPageShell`.",
        "points": [
            "Fetches leaderboard data from backend and combines it with current auth user from Redux.",
            "Highlights current user rank if present."
        ],
        "notes": ["Page uses both fetched public data and global auth data together."],
        "anchors": [
            ('export default function LeaderboardPage() {', "Defines leaderboard page."),
            ('const [leaders, setLeaders] = useState([]);', "Stores fetched leaderboard data."),
            ('useEffect(() => {', "Fetches leaderboard from backend."),
            ('const podium = leaders.slice(0, 3);', "Computes top three cards."),
            ('const currentUserRank = useMemo(', "Finds current user's rank."),
            ('return (', "Renders leaderboard layout."),
            ('function Metric({ label, value }) {', "Defines summary metric card."),
            ('function PodiumCard({ entry, highlight }) {', "Defines podium card."),
            ('function PodiumStat({ label, value }) {', "Defines small podium stat card."),
        ],
    },
    {
        "path": "src/Pages/LoginPage.jsx",
        "title": "Login Page",
        "purpose": "Implements user login form UI and client-side submission flow.",
        "flow": "Rendered at `/login` inside `LoginLayout`.",
        "points": [
            "Uses local form state and a submit trigger flag instead of form libraries.",
            "Sends `{ role: 'user' }` to backend login endpoint.",
            "Navigates to home on successful login."
        ],
        "notes": ["Password visibility toggle is implemented locally in the page."],
        "anchors": [
            ('export default function LoginPage() {', "Defines login page."),
            ('const [formFields, setFormFields] = useState({ email: "", password: "" });', "Stores form values."),
            ('useEffect(() => {', "Runs submit-side effect when `submit` flag flips."),
            ('await axios.post(', "Calls backend login endpoint."),
            ('setFormError(', "Stores API or validation error message."),
            ('return (', "Renders login form UI."),
        ],
    },
    {
        "path": "src/Pages/NotFoundPage.jsx",
        "title": "Not Found Page",
        "purpose": "Simple fallback 404 screen component.",
        "flow": "Intended for unmatched routes, but not currently wired into the router tree.",
        "points": ["Good interview point: component exists, but router integration is missing."],
        "notes": [],
        "anchors": [
            ('export default function NotFoundPage() {', "Defines the 404 page component."),
            ('<h1 className="text-3xl">404 - Page Not Found</h1>', "Renders the 404 message."),
        ],
    },
    {
        "path": "src/Pages/ProblemSetPage.jsx",
        "title": "Problem Set Page",
        "purpose": "Displays the full question bank with search, band filters, stats, and direct navigation to each problem.",
        "flow": "Rendered at `/problems` inside `ProblemLayout`; data comes from a route loader.",
        "points": [
            "Uses React Router loader for data fetch instead of `useEffect` inside component.",
            "Computes filtered problem view with `useMemo`."
        ],
        "notes": ["Design mirrors Explore page but is more list-oriented and directly task-focused."],
        "anchors": [
            ('function ProblemSetPage() {', "Defines the problem-set page."),
            ('const { problems } = useLoaderData();', "Receives preloaded data from loader."),
            ('const [search, setSearch] = useState("");', "Stores filter UI state."),
            ('const filteredProblems = useMemo(() => {', "Computes filtered list."),
            ('const stats = {', "Computes summary metrics."),
            ('return (', "Renders problem-set layout."),
            ('function Metric({ label, value }) {', "Defines inline stat card."),
            ('export const problemSetLoader = async () => {', "Defines route loader for problem set data."),
        ],
    },
    {
        "path": "src/Pages/ProblemPage.jsx",
        "title": "Problem Workspace Page",
        "purpose": "Core coding workspace page that shows the problem statement, code editor, console, AI review, and submission history.",
        "flow": "Rendered at `/problems/:problemId` inside `ProblemLayout`; data comes from a route loader and user actions trigger both backend and compiler-service calls.",
        "points": [
            "This is the strongest file to discuss in interviews because it combines layout, loaders, local state, editor integration, split panes, markdown rendering, backend calls, compiler calls, and submission history.",
            "Uses `react-split` for resizable panels and Monaco for code editing.",
            "Calls compiler service for run/review and backend for submit/submission history."
        ],
        "notes": [
            "The language effect contains a Java template, but Java is not exposed in the select options.",
            "Problem loader redirects to `/login` on backend 401."
        ],
        "anchors": [
            ('function ProblemPage() {', "Defines the main coding workspace page."),
            ('const { problem } = useLoaderData();', "Receives preloaded problem data."),
            ('const [code, setCode] = useState(', "Stores editor code state and other local workspace state."),
            ('const fetchSubs = async () => {', "Defines lazy submission-history fetch."),
            ('const toggleSubs = () => {', "Opens/closes submission history panel."),
            ('const handleRun = async () => {', "Runs code through compiler service `/run`."),
            ('const handleSubmit = async () => {', "Submits code to backend judge endpoint."),
            ('const handleReview = async () => {', "Requests AI review from compiler service."),
            ('useEffect(() => {', "Resets starter template when language changes."),
            ('return (', "Renders the full workspace UI."),
            ('function TopStat({ label, value }) {', "Defines header stat card."),
            ('export const ProblemLoader = async ({ params }) => {', "Defines route loader for problem detail."),
        ],
    },
    {
        "path": "src/Pages/ProfilePage.jsx",
        "title": "Profile Page",
        "purpose": "Shows the logged-in user's profile summary, stats, and submission history.",
        "flow": "Rendered at `/profile` inside `RootLayout`; page fetches profile and submissions directly on mount.",
        "points": [
            "Uses local state rather than Redux for profile details and submission history.",
            "Logout is driven through Redux thunk, but data loading is page-local."
        ],
        "notes": ["Short variable names like `u`, `su`, `s`, `ss`, `d`, and `n` make the implementation less readable than the rest of the codebase."],
        "anchors": [
            ('export default function ProfilePage() {', "Defines profile page."),
            ('const [u, su] = useState(null);', "Stores user profile locally."),
            ('const logout = async () => {', "Logs out through Redux and redirects."),
            ('useEffect(() => {', "Fetches profile and submissions on mount."),
            ('const acceptanceRate = useMemo(() => {', "Computes acceptance rate from loaded profile data."),
            ('if (!u) {', "Shows loading state until profile arrives."),
            ('return (', "Renders profile layout."),
            ('function Panel({ title, children }) {', "Defines reusable side panel."),
            ('function InfoRow({ label, value }) {', "Defines metadata row."),
        ],
    },
    {
        "path": "src/Pages/RegisterPage.jsx",
        "title": "Register Page",
        "purpose": "Implements user registration UI with local password-strength validation and backend submission.",
        "flow": "Rendered at `/register` inside `LoginLayout`.",
        "points": [
            "Runs client-side password complexity check before calling backend.",
            "Sends `{ role: 'user' }` to backend register endpoint."
        ],
        "notes": ["Backend still re-validates all fields, which is correct defense in depth."],
        "anchors": [
            ('export default function RegisterPage() {', "Defines register page."),
            ('const [formFields, setFormFields] = useState({ userName: "", email: "", password: "" });', "Stores registration form data."),
            ('const passwordRequirements =', "Defines visible password guidance text."),
            ('const isStrongPassword = (password) =>', "Defines client-side password rule."),
            ('useEffect(() => {', "Runs submit-side effect when submit flag changes."),
            ('await axios.post(', "Calls backend registration endpoint."),
            ('return (', "Renders registration form UI."),
        ],
    },
])

GUIDES.extend([
    {
        "path": "src/Components/AuthNavbar.jsx",
        "title": "Auth Navbar",
        "purpose": "Minimal navbar used on login/register screens.",
        "flow": "Rendered by `LoginLayout` above auth pages.",
        "points": ["Keeps auth screens visually simple and separate from the full main navbar."],
        "notes": [],
        "anchors": [
            ('export default function AuthNavbar() {', "Defines auth-specific navbar."),
            ('<nav className=', "Creates the top gradient navigation shell."),
            ('<Link to="/" className="flex items-center gap-3">', "Provides brand link to home."),
            ('to="/"', "Provides a direct way back to home."),
        ],
    },
    {
        "path": "src/Components/Footer.jsx",
        "title": "Footer",
        "purpose": "Shared footer shown on layouts that include it.",
        "flow": "Rendered by `RootLayout` and `LoginLayout`.",
        "points": ["Uses simple external links and dynamic current year."],
        "notes": [],
        "anchors": [
            ('export default function Footer() {', "Defines the shared footer component."),
            ('<footer className=', "Creates the footer container."),
            ('new Date().getFullYear()', "Computes current year dynamically."),
            ('href="https://github.com/Harsh3185/Code-Bharat"', "Links to project GitHub."),
        ],
    },
    {
        "path": "src/Components/MeshPageShell.jsx",
        "title": "Mesh Page Shell",
        "purpose": "Reusable page shell that applies the mesh background and atmospheric overlays.",
        "flow": "Used by several public pages like Explore, Contests, Leaderboard, Problem Set, and Profile.",
        "points": [
            "Centralizes the signature visual background so pages stay consistent.",
            "Accepts `contentClassName` for small layout variation without duplicating shell code."
        ],
        "notes": [],
        "anchors": [
            ('export default function MeshPageShell({ children, contentClassName = "" }) {', "Defines the reusable shell component."),
            ('<main className=', "Creates the background canvas."),
            ('<img', "Places the mesh image as full-page background."),
            ('<section className={`relative', "Wraps page content in a centered max-width container."),
        ],
    },
    {
        "path": "src/Components/Navbar.jsx",
        "title": "Main Navbar",
        "purpose": "Primary navigation for the public app, with auth-aware rendering and responsive mobile menu behavior.",
        "flow": "Rendered by `RootLayout` and `ProblemLayout` before page content.",
        "points": [
            "Hydrates auth state by dispatching `fetchUser` when needed.",
            "Shows admin shortcut only for admin users.",
            "Desktop and mobile nav behavior are both handled in one component."
        ],
        "notes": [
            "Navbar is a practical place where global Redux auth state is consumed.",
            "Because RootLayout renders Navbar before HomePage, auth state often hydrates there first."
        ],
        "anchors": [
            ('export default function Navbar() {', "Defines the main navbar component."),
            ('const [open, setOpen] = useState(false);', "Stores mobile menu open state."),
            ('const { pathname } = useLocation();', "Uses current URL to highlight active nav items."),
            ('const { user } = useSelector((s) => s.auth);', "Reads auth state from Redux."),
            ('useEffect(() => {', "Fetches current user if auth state is empty."),
            ('const nav = [', "Defines top-level navigation links."),
            ('return (', "Starts rendering navbar UI."),
            ('{user ? (', "Branches UI based on logged-in vs logged-out state."),
            ('<button', "Defines mobile hamburger toggle."),
        ],
    },
    {
        "path": "src/Components/SubmissionTable.jsx",
        "title": "Submission Table",
        "purpose": "Shared table for displaying submission history.",
        "flow": "Used on profile page and problem page submission history panel.",
        "points": [
            "Centralizing the table avoids repeating status rendering and date formatting logic.",
            "Links submission rows back to the problem detail page."
        ],
        "notes": [],
        "anchors": [
            ('export default function SubmissionTable({ submissions = [] }) {', "Defines shared submission table component."),
            ('if (!submissions.length) return null;', "Skips rendering when there is no data."),
            ('<table className=', "Builds table markup."),
            ('{submissions.map((submission, index) => (', "Renders one row per submission."),
            ('function statusClasses(status) {', "Maps statuses to badge color classes."),
        ],
    },
    {
        "path": "src/Layouts/AdminAuthLayout.jsx",
        "title": "Admin Auth Layout",
        "purpose": "Simple wrapper layout for admin login and admin registration screens.",
        "flow": "Applied to `/admin/login` and `/admin/register` routes.",
        "points": ["Keeps admin auth pages separate from the public root/login layouts."],
        "notes": [],
        "anchors": [
            ('export default function AdminAuthLayout() {', "Defines the admin auth layout."),
            ('<Outlet />', "Renders the nested admin auth page."),
        ],
    },
    {
        "path": "src/Layouts/AdminLayout.jsx",
        "title": "Admin Layout",
        "purpose": "Provides the protected admin workspace shell with sidebar navigation and logout action.",
        "flow": "Wraps admin dashboard and admin problem-management pages after `RequireAdmin` passes.",
        "points": [
            "Sidebar layout clearly separates admin workspace from user-facing UI.",
            "Logout is driven through Redux thunk and redirects to admin login."
        ],
        "notes": ["The `Open main site` link uses plain `<a href=\"/\">`, which triggers a full page reload instead of client-side navigation."],
        "anchors": [
            ('const links = [', "Defines sidebar navigation items."),
            ('export default function AdminLayout() {', "Defines the admin shell."),
            ('const { user } = useSelector((state) => state.auth);', "Reads current admin user from Redux."),
            ('const handleLogout = async () => {', "Logs out and redirects."),
            ('<aside className=', "Creates the admin sidebar."),
            ('<NavLink', "Renders active-aware sidebar navigation."),
            ('<Outlet />', "Renders nested admin page content."),
        ],
    },
    {
        "path": "src/Layouts/LoginLayout.jsx",
        "title": "Login Layout",
        "purpose": "Wraps user login and register pages with auth navbar and footer.",
        "flow": "Applied to `/login` and `/register` routes.",
        "points": ["Auth pages reuse the same top and bottom chrome while keeping forms page-specific."],
        "notes": [],
        "anchors": [
            ('function LoginLayout() {', "Defines login/register layout."),
            ('<AuthNavbar />', "Places simplified auth navbar above form pages."),
            ('<Outlet />', "Renders nested auth page."),
            ('<Footer />', "Adds footer below auth page."),
        ],
    },
    {
        "path": "src/Layouts/ProblemLayout.jsx",
        "title": "Problem Layout",
        "purpose": "Wraps the problem list and coding workspace with the main navbar.",
        "flow": "Applied to `/problems` and `/problems/:problemId`.",
        "points": ["Problem area has its own layout so workspace pages keep a consistent frame without the footer."],
        "notes": [],
        "anchors": [
            ('function ProblemLayout() {', "Defines problem-area layout."),
            ('<Navbar />', "Renders main navbar."),
            ('<Outlet />', "Renders nested problem page content."),
        ],
    },
    {
        "path": "src/Layouts/RootLayout.jsx",
        "title": "Root Layout",
        "purpose": "Main public-site layout with navbar, routed content, and footer.",
        "flow": "Applied to home, profile, contests, leaderboard, and explore routes.",
        "points": ["Provides consistent shell for the public site."],
        "notes": [],
        "anchors": [
            ('function RootLayout() {', "Defines main public layout."),
            ('<Navbar />', "Renders shared top navigation."),
            ('<Outlet />', "Renders nested public page."),
            ('<Footer />', "Renders shared footer."),
        ],
    },
    {
        "path": "src/Middleware/RequireAdmin.jsx",
        "title": "Require Admin Guard",
        "purpose": "Protects admin routes by fetching current user and redirecting non-admins.",
        "flow": "Runs before `AdminLayout` and all nested admin pages.",
        "points": [
            "Acts like client-side route middleware using React Router `<Navigate />`.",
            "Relies on Redux auth state and `fetchUser` thunk."
        ],
        "notes": ["This is frontend route protection; real security still depends on backend auth and admin checks."],
        "anchors": [
            ('export default function RequireAdmin() {', "Defines the admin route guard."),
            ('const { user, loading } = useSelector((state) => state.auth);', "Reads auth state from Redux."),
            ('useEffect(() => {', "Fetches current user if not already loaded."),
            ('if (loading) {', "Shows loading screen while auth is resolving."),
            ('if (!user || user.role?.toLowerCase() !== "admin") {', "Redirects non-admin users."),
            ('return <Outlet />;', "Allows admin route tree to render."),
        ],
    },
])


ROUTES = [
    ("/", "RootLayout", "HomePage", "Public landing page"),
    ("/profile", "RootLayout", "ProfilePage", "Authenticated profile experience handled in page logic"),
    ("/contests", "RootLayout", "ContestsPage", "Contest packs built from live problems"),
    ("/leaderboard", "RootLayout", "LeaderboardPage", "Public ranking page"),
    ("/explore", "RootLayout", "ExplorePage", "Searchable discovery page"),
    ("/problems", "ProblemLayout", "ProblemSetPage", "Problem list with loader"),
    ("/problems/:problemId", "ProblemLayout", "ProblemPage", "Problem workspace with loader"),
    ("/login", "LoginLayout", "LoginPage", "User login"),
    ("/register", "LoginLayout", "RegisterPage", "User registration"),
    ("/admin/login", "AdminAuthLayout", "AdminLoginPage", "Admin login"),
    ("/admin/register", "AdminAuthLayout", "AdminRegisterPage", "Admin registration"),
    ("/admin", "RequireAdmin -> AdminLayout", "AdminDashboardPage", "Protected admin dashboard"),
    ("/admin/problems", "RequireAdmin -> AdminLayout", "AdminProblemsPage", "Protected problem management"),
    ("/admin/problems/new", "RequireAdmin -> AdminLayout", "AdminAddProblemPage", "Protected problem creation"),
]


ENV_ROWS = [
    ("VITE_BACKEND_URL_DEV", "Development backend base URL override"),
    ("VITE_COMPILER_URL_DEV", "Development compiler-service URL override"),
    ("VITE_BACKEND_URL", "Production backend base URL"),
    ("VITE_COMPILER_URL", "Production compiler-service base URL"),
]


COMPONENT_GROUPS = [
    ("Entry", "index.html, main.jsx, App.jsx"),
    ("Config", "vite.config.js, eslint.config.js, config/urls.js"),
    ("State", "Redux store, auth slice, auth thunks, auth API"),
    ("Layouts", "RootLayout, ProblemLayout, LoginLayout, AdminAuthLayout, AdminLayout"),
    ("Middleware", "RequireAdmin route guard"),
    ("Components", "Navbar, AuthNavbar, Footer, MeshPageShell, SubmissionTable"),
    ("Pages", "Home, Explore, Contests, Leaderboard, ProblemSet, Problem, Profile, Login, Register, admin pages"),
]


def add_front(doc):
    title(doc, "Code Bharat Frontend Interview Guide")
    title(doc, f"Generated from this repository on {datetime.now().strftime('%Y-%m-%d %H:%M')}", 11, True)
    doc.add_paragraph("Goal: explain how the frontend is structured, how each page is designed and implemented, how routing/state work, and how UI code talks to the backend.")
    doc.add_heading("1. Frontend Overview", level=1)
    doc.add_paragraph("This frontend is a React 19 + Vite single-page application styled with Tailwind CSS utilities, animated with Framer Motion, state-managed with Redux Toolkit for auth, and routed with React Router. It has a public experience, an authenticated problem-solving workspace, and a separate admin workspace.")
    doc.add_heading("2. Render Flow", level=1)
    code(doc, "\n".join([
        "index.html",
        "  -> src/main.jsx",
        "    -> Redux Provider(store)",
        "      -> App.jsx",
        "        -> React Router tree",
        "          -> Layout",
        "            -> Page",
        "              -> Shared components",
        "                -> axios calls to backend/compiler",
    ]))
    doc.add_heading("3. How UI State Travels", level=1)
    nums(doc, [
        "App boot starts in `main.jsx`, which wraps the whole tree in Redux `Provider`.",
        "`App.jsx` builds the browser router and selects layout + page by URL.",
        "Layouts provide shared frame elements like navigation, footer, admin sidebar, or page chrome.",
        "Pages own most of the screen-specific fetching and local UI state.",
        "Shared auth state lives in Redux and is fetched through `fetchUser` thunk.",
        "Axios calls use URLs from `src/config/urls.js` and send cookies with `withCredentials: true`.",
        "Problem solving also calls the separate compiler service for run/review actions.",
    ])
    doc.add_heading("4. Route Map", level=1)
    table(doc, ["Path", "Layout / Guard", "Page", "Purpose"], ROUTES)
    doc.add_heading("5. File Groups", level=1)
    table(doc, ["Area", "Files"], COMPONENT_GROUPS)
    doc.add_heading("6. Environment Variables", level=1)
    table(doc, ["Variable", "Use"], ENV_ROWS)
    doc.add_heading("7. Source Tree", level=1)
    code(doc, tree())
    doc.add_heading("8. Key Interview Observations", level=1)
    bullets(doc, [
        "The app uses Redux only for auth; most page state is local `useState` + `useEffect`.",
        "Data fetching is distributed across pages instead of centralized in a service layer or React Query.",
        "The router imports `NotFoundPage`, but there is no catch-all route defined, so the 404 component is currently unused.",
        "`useAuth` hook exists, but navbar/admin guard rely on Redux thunk fetch instead, so there are two auth-access patterns.",
        "`App.css` exists but is effectively unused; `index.css` contains the active global styles.",
        "Problem workspace mixes backend submission calls and compiler-service run/review calls in a single page component.",
        "Admin problem creation is a two-step workflow: create problem first, then create hidden test cases.",
    ])


def add_env_page(doc):
    env = FRONTEND / ".env"
    if not env.exists():
        return
    redacted = []
    for line in read(env).splitlines():
        if not line.strip() or line.lstrip().startswith("#"):
            redacted.append(line)
        else:
            redacted.append(line.split("=", 1)[0] + "=***REDACTED***")
    doc.add_page_break()
    doc.add_heading("Redacted Frontend `.env`", level=1)
    code(doc, numbered("\n".join(redacted)))


def add_file(doc, entry):
    text = read(FRONTEND / entry["path"])
    doc.add_page_break()
    doc.add_heading(entry["title"], level=1)
    doc.add_paragraph(f"File: Frontend/{entry['path']}")
    doc.add_paragraph(f"Purpose: {entry['purpose']}")
    doc.add_paragraph(f"How it fits into the app: {entry['flow']}")
    doc.add_paragraph("Interview-ready talking points:")
    bullets(doc, entry["points"])
    if entry.get("notes"):
        doc.add_paragraph("Implementation notes:")
        bullets(doc, entry["notes"])
    doc.add_heading("Code", level=2)
    code(doc, numbered(text))
    doc.add_heading("Walkthrough", level=2)
    walkthrough(doc, text, entry["anchors"])


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style(doc)
    add_front(doc)
    add_env_page(doc)
    for entry in GUIDES:
        add_file(doc, entry)
    doc.add_page_break()
    doc.add_heading("How To Explain This Frontend In Interviews", level=1)
    bullets(doc, [
        "Start with the stack: React + Vite + Tailwind + Redux Toolkit + React Router + axios.",
        "Explain route hierarchy from `App.jsx`, then explain what each layout contributes.",
        "Describe auth as cookie-based and shared through Redux auth state.",
        "Explain that most UI pages fetch their own data with axios and local state.",
        "Use `ProblemPage.jsx` as the strongest example because it combines loaders, editor UI, compiler calls, submission calls, markdown rendering, split panes, and submission history.",
        "Use admin pages as an example of separating privileged flows from normal user flows.",
    ])
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()

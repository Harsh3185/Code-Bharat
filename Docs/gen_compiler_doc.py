from pathlib import Path
from datetime import datetime
from docx import Document
from gen_doc import style, title, bullets, nums, code, numbered, table, walkthrough, read

ROOT = Path(__file__).resolve().parent
COMPILER = ROOT / "Compiler"
OUT = ROOT / "output" / "doc" / "Code-Bharat-Compiler-Interview-Guide.docx"


def tree():
    files = sorted(
        p.relative_to(COMPILER).as_posix()
        for p in COMPILER.rglob("*")
        if p.is_file()
        and "node_modules" not in p.parts
        and p.parts[-2:-1] != ("codes",)
        and p.parts[-2:-1] != ("inputs",)
        and p.parts[-2:-1] != ("outputs",)
        and "Keys" not in p.parts
        and p.name != "package-lock.json"
    )
    out = ["Compiler/"]
    for f in files:
        out.append(f"{'  ' * f.count('/')} - {f.split('/')[-1]}")
    return "\n".join(out)


GUIDES = [
    {
        "path": "package.json",
        "title": "Compiler Package Manifest",
        "purpose": "Defines the compiler service runtime, scripts, and dependencies.",
        "flow": "This file shapes how the compiler service starts and which libraries are available to `/run` and `/ai-review`.",
        "points": [
            "The service is a Node CommonJS app started from `Index.js`.",
            "Dependencies are intentionally small: Express, CORS, dotenv, UUID, and Google GenAI."
        ],
        "notes": ["Package name is still `backend`, which is misleading for a separate compiler service."],
        "anchors": [
            ('"scripts": {', "Defines service start commands."),
            ('"type": "commonjs"', "Declares CommonJS module mode."),
            ('"dependencies": {', "Lists runtime dependencies used by the compiler service."),
        ],
    },
    {
        "path": ".gitignore",
        "title": "Compiler Git Ignore Rules",
        "purpose": "Prevents secrets, generated temp files, and local artifacts from being committed.",
        "flow": "Does not affect request handling but protects repo hygiene.",
        "points": [
            "Generated source/input/output temp folders are ignored.",
            "Secrets and private keys are intended to be ignored."
        ],
        "notes": ["Ignore rule uses `keys/` while repo folder is `Keys/`; Windows is case-insensitive, but other environments may not be."],
        "anchors": [
            ('.env', "Ignores secret environment file."),
            ('keys/', "Intends to ignore private key folder."),
            ('node_modules/', "Ignores installed dependencies."),
            ('codes/', "Ignores generated source files."),
            ('outputs/', "Ignores compiled binaries."),
            ('inputs/', "Ignores generated input files."),
        ],
    },
    {
        "path": ".dockerignore",
        "title": "Compiler Docker Ignore Rules",
        "purpose": "Keeps Docker build context small by excluding local and generated artifacts.",
        "flow": "Affects container build behavior only.",
        "points": ["Generated temp folders are excluded from Docker context."],
        "notes": ["Lockfile is ignored here."],
        "anchors": [
            ('node_modules', "Skips local dependencies."),
            ('package-lock.json', "Skips package lockfile in Docker context."),
            ('codes', "Skips generated code files."),
            ('outputs', "Skips compiled binaries."),
            ('inputs', "Skips generated input files."),
        ],
    },
    {
        "path": "Dockerfile",
        "title": "Compiler Dockerfile",
        "purpose": "Builds a runnable container with native toolchains for compiling and executing user code.",
        "flow": "Defines the environment that the compiler service needs in production: C/C++ build tools, Python, Go, and JDK.",
        "points": [
            "Installs native toolchains because `/run` shells out to system compilers/interpreters.",
            "Exposes port 7000, matching the service default."
        ],
        "notes": ["Docker uses Node 18, but `package.json` engines says 20.x. That mismatch is worth discussing in interviews."],
        "anchors": [
            ('FROM node:18', "Selects the base Node image."),
            ('RUN apt-get update && apt-get install -y \\', "Installs required native toolchains."),
            ('COPY .env .env', "Copies environment file into image."),
            ('COPY package*.json ./', "Copies package metadata before install."),
            ('RUN npm install', "Installs service dependencies."),
            ('COPY . .', "Copies service source code."),
            ('EXPOSE 7000', "Documents service port."),
            ('CMD ["node", "Index.js"]', "Starts compiler service."),
        ],
    },
    {
        "path": "Index.js",
        "title": "Compiler Service Entrypoint",
        "purpose": "Creates the Express service, exposes `/run` and `/ai-review`, writes temp files, dispatches execution by language, and returns results.",
        "flow": "Backend `submitSolution` and frontend `ProblemPage` both depend on this service for code execution and AI review.",
        "points": [
            "This file is the orchestration layer for compiler operations.",
            "Run flow goes request -> validate code -> generate temp code/input files -> select executor -> return output.",
            "AI review flow goes request -> Gemini call -> markdown/text review response."
        ],
        "notes": [
            "`executeJava` is implemented in the repo but not wired into the `/run` language switch.",
            "Returned response includes file paths, which is useful for debugging but may expose internal temp-path details."
        ],
        "anchors": [
            ("const express = require('express');", "Imports Express and all compiler helper modules."),
            ('app.use(cors());', "Enables cross-origin requests."),
            ('app.get("/", (req, res) => {', "Defines a health endpoint."),
            ('app.post("/run", async (req, res) => {', "Defines the main code-execution endpoint."),
            ('const { language = \'cpp\', code, input = \'\' } = req.body;', "Reads run request payload."),
            ('if (code === undefined || code.trim() === \'\') {', "Rejects empty code submissions early."),
            ('const filePath = await generateFile(language, code);', "Writes source code to disk."),
            ('const inputPath = await generateInputFile(input);', "Writes stdin data to disk."),
            ('if (language === "cpp") {', "Selects execution helper by language."),
            ('res.json({', "Returns output and temp file metadata."),
            ('app.post("/ai-review", async (req, res) => {', "Defines AI review endpoint."),
            ('const review = await aiCodeReview(code);', "Calls Gemini review helper."),
            ('const PORT = process.env.PORT || 7000;', "Starts server on configured port."),
        ],
    },
    {
        "path": "aiCodeReview.js",
        "title": "AI Review Helper",
        "purpose": "Wraps Google Gemini content generation to review submitted code.",
        "flow": "Called only by `/ai-review` in `Index.js`.",
        "points": [
            "Loads Gemini API key from environment.",
            "Uses `gemini-2.5-flash` with a direct prompt that embeds the user's code."
        ],
        "notes": ["There is no prompt guard, usage limit, or output post-processing layer."],
        "anchors": [
            ('const dotenv = require("dotenv");', "Imports dotenv and Gemini SDK."),
            ('dotenv.config();', "Loads Gemini API key from environment."),
            ('const ai = new GoogleGenAI({ apiKey: process.env.GOOGLE_GEMINI_API_KEY });', "Creates Gemini client."),
            ('const aiCodeReview = async (code) => {', "Defines review helper."),
            ('const response = await ai.models.generateContent({', "Calls Gemini model."),
            ('return response.text;', "Returns generated review text."),
        ],
    },
    {
        "path": "generateFile.js",
        "title": "Source File Generator",
        "purpose": "Writes submitted code to a temp file in the `codes/` directory.",
        "flow": "Called before execution so native tools can compile or interpret a real file.",
        "points": [
            "Uses UUIDs for most filenames.",
            "Java is special-cased to `Main.java` so class name matches executor expectations."
        ],
        "notes": ["Fixed `Main.java` naming can create collisions if multiple Java jobs run concurrently."],
        "anchors": [
            ("const fs = require('fs');", "Imports filesystem, path, and UUID helpers."),
            ("const dirCodes = path.join(__dirname, 'codes');", "Defines temp source directory."),
            ("if (!fs.existsSync(dirCodes)) {", "Creates source directory if needed."),
            ('const generateFile = (format, content) => {', "Defines source-file generator."),
            ('const jobID = uuid();', "Creates unique job id."),
            ('const filename = format === "java" ? "Main.java" : `${jobID}.${format}`;', "Selects filename pattern."),
            ('fs.writeFileSync(filePath, content);', "Writes source code to disk."),
        ],
    },
    {
        "path": "generateInputFile.js",
        "title": "Input File Generator",
        "purpose": "Writes stdin input to a temp `.txt` file in the `inputs/` directory.",
        "flow": "Called before execution so command-line programs can read redirected stdin.",
        "points": ["Uses UUID-based filenames to isolate jobs."],
        "notes": [],
        "anchors": [
            ("const fs = require('fs');", "Imports filesystem, path, and UUID helpers."),
            ("const dirInputs = path.join(__dirname, 'inputs');", "Defines temp input directory."),
            ("if (!fs.existsSync(dirInputs)) {", "Creates input directory if needed."),
            ('const generateInputFile = async (input) => {', "Defines input-file generator."),
            ('const jobID = uuid();', "Creates unique input file id."),
            ('await fs.writeFileSync(input_filePath, input);', "Writes input text to disk."),
        ],
    },
    {
        "path": "executeC.js",
        "title": "C Executor",
        "purpose": "Compiles C code with `gcc`, executes the compiled binary, and cleans temp files.",
        "flow": "Called by `/run` when language is `c`.",
        "points": ["Run flow is compile command plus stdin redirection command combined with `&&`."],
        "notes": ["Compilation and runtime errors are returned together via rejected object."],
        "anchors": [
            ('const executeC = (filepath, inputPath) => {', "Defines C execution helper."),
            ('const dir = path.dirname(filepath);', "Computes output location."),
            ('const compileCmd = `gcc', "Builds compile command."),
            ('const runCmd = `"${outPath}" < "${inputPath}"`;', "Builds run command with stdin redirection."),
            ('exec(`${compileCmd} && ${runCmd}`', "Runs compile then execute."),
            ('cleanupFiles([filepath, inputPath, outPath]);', "Deletes temp files after execution."),
        ],
    },
    {
        "path": "executeCpp.js",
        "title": "C++ Executor",
        "purpose": "Compiles C++ code with `g++`, executes the binary, and cleans temp files.",
        "flow": "Called by `/run` when language is `cpp`.",
        "points": ["Compiled `.exe` is stored under `outputs/` before cleanup."],
        "notes": [],
        "anchors": [
            ('const outputPath = path.join(__dirname, "outputs");', "Defines compiled-binary output directory."),
            ('if (!fs.existsSync(outputPath)) fs.mkdirSync(outputPath, { recursive: true });', "Creates binary output directory."),
            ('const executeCpp = (filepath, inputPath) => {', "Defines C++ execution helper."),
            ('const compileCmd = `g++', "Builds compile command."),
            ('exec(`${compileCmd} && ${runCmd}`', "Runs compile then execute."),
            ('cleanupFiles([filepath, inputPath, outPath]);', "Deletes temp files after completion."),
        ],
    },
    {
        "path": "executeGo.js",
        "title": "Go Executor",
        "purpose": "Runs Go source directly with `go run` and cleans temp files.",
        "flow": "Called by `/run` when language is `go`.",
        "points": ["Go path skips a separate compiled-output file and uses `go run` directly."],
        "notes": [],
        "anchors": [
            ('const executeGo = (filepath, inputPath) => {', "Defines Go execution helper."),
            ('exec(`go run "${filepath}" < "${inputPath}"`', "Runs Go code with redirected stdin."),
            ('cleanupFiles([filepath, inputPath]);', "Deletes temp source and input files."),
        ],
    },
    {
        "path": "executeJava.js",
        "title": "Java Executor",
        "purpose": "Compiles `Main.java`, runs the resulting class, and cleans generated files.",
        "flow": "Implemented but not currently wired into `/run` in `Index.js`.",
        "points": ["Java compile/run flow is present and interview-worth noting even though it is currently unreachable from the route switch."],
        "notes": ["If Java is re-enabled in UI and service switch, this executor is ready for that path."],
        "anchors": [
            ('const executeJava = (filepath, inputPath) => {', "Defines Java execution helper."),
            ('const className = "Main";', "Uses fixed class name expected by generated file."),
            ('const compileCmd = `javac', "Builds Java compile command."),
            ('const runCmd = `java -cp', "Builds Java run command."),
            ('exec(`${compileCmd} && ${runCmd}`', "Runs compile then execute."),
            ('cleanupFiles([', "Deletes `.java`, input, and compiled `.class` file."),
        ],
    },
    {
        "path": "executeJs.js",
        "title": "JavaScript Executor",
        "purpose": "Runs JavaScript code with Node and cleans temp files.",
        "flow": "Called by `/run` when language is `javascript`.",
        "points": ["Uses stdin redirection into `node` process."],
        "notes": [],
        "anchors": [
            ('const executeJs = (filepath, inputPath) => {', "Defines JavaScript execution helper."),
            ('exec(`node "${filepath}" < "${inputPath}"`', "Runs Node against temp script file."),
            ('cleanupFiles([filepath, inputPath]);', "Deletes temp source and input files."),
        ],
    },
    {
        "path": "executePython.js",
        "title": "Python Executor",
        "purpose": "Runs Python code with `python3` and cleans temp files.",
        "flow": "Called by `/run` when language is `python`.",
        "points": ["Uses interpreter execution instead of compiling a binary first."],
        "notes": ["The command assumes `python3` exists in PATH, which is true in the Docker image but can differ on Windows host setups."],
        "anchors": [
            ('const executePython = (filepath, inputPath) => {', "Defines Python execution helper."),
            ('exec(`python3 "${filepath}" < "${inputPath}"`', "Runs Python script with redirected stdin."),
            ('cleanupFiles([filepath, inputPath]);', "Deletes temp source and input files."),
        ],
    },
    {
        "path": "utils/cleanup.js",
        "title": "Cleanup Utility",
        "purpose": "Deletes temp files asynchronously after execution completes.",
        "flow": "Called by every executor to remove generated source/input/output artifacts.",
        "points": ["Keeps temp-file growth under control after each job."],
        "notes": ["Cleanup is best-effort and logs failures except missing-file cases."],
        "anchors": [
            ('const cleanupFiles = (paths = []) => {', "Defines cleanup helper."),
            ('for (const p of paths) {', "Loops over temp file paths."),
            ('fs.unlink(p, (err) => {', "Deletes each file asynchronously."),
            ("if (err && err.code !== 'ENOENT') {", "Logs real cleanup failures but ignores already-missing files."),
        ],
    },
]


ENV_ROWS = [
    ("PORT", "Compiler service listen port"),
    ("GOOGLE_GEMINI_API_KEY", "Gemini API key used by `/ai-review`"),
]


def add_front(doc):
    title(doc, "Code Bharat Compiler Interview Guide")
    title(doc, f"Generated from this repository on {datetime.now().strftime('%Y-%m-%d %H:%M')}", 11, True)
    doc.add_paragraph("Goal: explain how the compiler service works, how code moves from request to execution, how temporary files are created and cleaned, and how the backend depends on this service.")
    doc.add_heading("1. Compiler Service Overview", level=1)
    doc.add_paragraph("This compiler is a separate Express microservice that receives code, writes temporary source and input files, runs language-specific compile/execute commands, returns output, and optionally asks Gemini for AI review.")
    doc.add_heading("2. Run Request Flow", level=1)
    nums(doc, [
        "Frontend or backend sends `POST /run` with `language`, `code`, and optional `input`.",
        "Service validates that code is not empty.",
        "Service writes source code to `codes/` and input text to `inputs/`.",
        "Service chooses executor by language.",
        "Executor shells out to native toolchain/interpreter with stdin redirection.",
        "Executor resolves stdout or rejects with error details.",
        "Cleanup helper removes temporary files.",
        "Service returns output JSON to caller.",
    ])
    doc.add_heading("3. AI Review Flow", level=1)
    nums(doc, [
        "Frontend sends `POST /ai-review` with `code`.",
        "Service validates code presence.",
        "Service calls Gemini through `aiCodeReview.js`.",
        "Generated review text is returned to the frontend.",
    ])
    doc.add_heading("4. Source Tree", level=1)
    code(doc, tree())
    doc.add_heading("5. Environment Variables", level=1)
    table(doc, ["Variable", "Use"], ENV_ROWS)
    doc.add_heading("6. Key Interview Observations", level=1)
    bullets(doc, [
        "This service is execution-orchestration heavy and depends on OS-level compilers/interpreters, so deployment environment matters.",
        "The backend submission flow depends on this service for every hidden test case.",
        "Java executor exists but `/run` does not currently dispatch to it.",
        "Fixed `Main.java` naming can cause concurrency issues for parallel Java jobs.",
        "There is no timeout, sandboxing, queueing, or resource isolation beyond basic process execution.",
        "Temp-file cleanup is present, which is good, but failures are only logged.",
    ])


def add_env_page(doc):
    env = COMPILER / ".env"
    if not env.exists():
        return
    redacted = []
    for line in read(env).splitlines():
        if not line.strip() or line.lstrip().startswith("#"):
            redacted.append(line)
        else:
            redacted.append(line.split("=", 1)[0] + "=***REDACTED***")
    doc.add_page_break()
    doc.add_heading("Redacted Compiler `.env`", level=1)
    code(doc, numbered("\n".join(redacted)))


def add_file(doc, entry):
    text = read(COMPILER / entry["path"])
    doc.add_page_break()
    doc.add_heading(entry["title"], level=1)
    doc.add_paragraph(f"File: Compiler/{entry['path']}")
    doc.add_paragraph(f"Purpose: {entry['purpose']}")
    doc.add_paragraph(f"How it fits into service flow: {entry['flow']}")
    doc.add_paragraph("Interview-ready talking points:")
    bullets(doc, entry["points"])
    if entry.get("notes"):
        doc.add_paragraph("Implementation notes:")
        bullets(doc, entry["notes"])
    doc.add_heading("Code", level=2)
    code(doc, numbered(text))
    doc.add_heading("Walkthrough", level=2)
    walkthrough(doc, text, entry["anchors"])


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style(doc)
    add_front(doc)
    add_env_page(doc)
    for entry in GUIDES:
        add_file(doc, entry)
    doc.add_page_break()
    doc.add_heading("How To Explain This Compiler In Interviews", level=1)
    bullets(doc, [
        "Describe it as a separate execution microservice used by the main backend judge flow.",
        "Explain the file lifecycle: incoming code/input -> temp files -> native toolchain -> stdout -> cleanup.",
        "Mention that each language has its own executor wrapper around shell commands.",
        "Call out the risks: no sandboxing, no timeouts, no queueing, and environment/toolchain dependency.",
        "Mention AI review as a separate Gemini-backed endpoint, not part of correctness judging.",
    ])
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
